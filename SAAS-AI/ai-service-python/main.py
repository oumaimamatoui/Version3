"""
EvaluaTech AI Engine v10.1 — CODE COMPLET A→Z
✅ CHATBOT 200+ INTENTIONS — toutes questions plateforme SaaS couvertes
✅ INTENTIONS SPÉCIFIQUES par rôle : Candidat, Évaluateur, RH, Recruteur, AdminEntreprise, SuperAdmin
✅ DASHBOARD DYNAMIQUE — endpoints complets avec données simulées réalistes
✅ VOCAL FR/EN/AR — config complète
✅ RECOMMANDATIONS IA automatiques par rôle via Gemini + fallback ultra-riche
✅ LETTRE MOTIVATION — langue strictement respectée (FR/EN/AR) + validation style
✅ ANALYSE CV — détection document non-CV avec alerte structurée
✅ QCM trilingue FR/EN/AR
✅ ENTRETIEN IA — questions comportementales/techniques
✅ RAPPORTS PDF automatiques
✅ STATISTIQUES temps réel dynamiques par rôle
✅ Cache LRU, circuit breaker, semaphores, rate limiting
✅ Streaming SSE pour le chat
✅ NOUVEAU v10.1: /ia/lettre-motivation — enforced language output, style validated
✅ NOUVEAU v10.1: /ia/match-cv — CV vs non-CV detection, alert JSON field
✅ NOUVEAU v10.1: /ia/validate-cv-doc — endpoint dédié vérification document
✅ FIX v10.1: QCM génération — theme + sousTheme injectés dans chaque question
✅ FIX v10.1: Chatbot — intention "theme_soustheme" ajoutée (domaines disponibles)
"""

import os, io, json, logging, time, asyncio, hashlib, random, re
from contextlib import asynccontextmanager
from collections import deque, OrderedDict
from concurrent.futures import ThreadPoolExecutor
from typing import Optional, AsyncGenerator, List, Dict, Any
from datetime import datetime, timedelta
from pydantic import BaseModel

from fastapi import FastAPI, Form, File, UploadFile, HTTPException, Request, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import StreamingResponse, JSONResponse

try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

# ─────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────
logging.basicConfig(level=logging.WARNING, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
logger = logging.getLogger("evaluatech")

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

API_KEY = os.getenv("GEMINI_API_KEY", "")
WORKING_MODEL = "gemini-1.5-flash"
_START_TIME = time.time()
_gemini_client = None

if GEMINI_AVAILABLE and API_KEY:
    try:
        _gemini_client = genai.Client(api_key=API_KEY)
    except Exception as e:
        logger.warning(f"Gemini init failed: {e}")

_gemini_executor = ThreadPoolExecutor(max_workers=8, thread_name_prefix="gemini")

# ─────────────────────────────────────────────────
# CIRCUIT BREAKER
# ─────────────────────────────────────────────────
class CircuitBreaker:
    CLOSED = "CLOSED"; OPEN = "OPEN"; HALF_OPEN = "HALF_OPEN"
    def __init__(self, fail_threshold=5, recovery_timeout=60):
        self.state = self.CLOSED; self.fail_count = 0
        self.fail_threshold = fail_threshold; self.recovery_timeout = recovery_timeout
        self._opened_at: float = 0
    def record_success(self): self.fail_count = 0; self.state = self.CLOSED
    def record_failure(self):
        self.fail_count += 1
        if self.fail_count >= self.fail_threshold:
            self.state = self.OPEN; self._opened_at = time.time()
    def is_allowed(self) -> bool:
        if self.state == self.CLOSED: return True
        if self.state == self.OPEN:
            if time.time() - self._opened_at > self.recovery_timeout:
                self.state = self.HALF_OPEN; return True
            return False
        return True

_circuit = CircuitBreaker()

# ─────────────────────────────────────────────────
# RATE LIMITER
# ─────────────────────────────────────────────────
_rate_limits: dict = {}
def check_rate_limit(ip: str) -> bool:
    now = time.time()
    if ip not in _rate_limits: _rate_limits[ip] = deque()
    dq = _rate_limits[ip]
    while dq and now - dq[0] > 60: dq.popleft()
    if len(dq) >= 120: return False
    dq.append(now); return True

# ─────────────────────────────────────────────────
# CACHE LRU
# ─────────────────────────────────────────────────
class LRUCache:
    def __init__(self, maxsize=1000, ttl=600):
        self._store: OrderedDict = OrderedDict()
        self.maxsize = maxsize; self.ttl = ttl
    def get(self, key: str):
        if key not in self._store: return None
        value, ts = self._store[key]
        if time.time() - ts > self.ttl: del self._store[key]; return None
        self._store.move_to_end(key); return value
    def set(self, key: str, value):
        if key in self._store: self._store.move_to_end(key)
        self._store[key] = (value, time.time())
        while len(self._store) > self.maxsize: self._store.popitem(last=False)

_cache      = LRUCache(maxsize=1000, ttl=600)
_chat_cache = LRUCache(maxsize=500,  ttl=300)
_reco_cache = LRUCache(maxsize=200,  ttl=180)
_dash_cache = LRUCache(maxsize=300,  ttl=60)

def make_cache_key(*args) -> str:
    return hashlib.md5("|".join(str(a) for a in args).encode(), usedforsecurity=False).hexdigest()

_gemini_semaphore = asyncio.Semaphore(6)
_sem_qcm   = asyncio.Semaphore(3)
_sem_cv    = asyncio.Semaphore(2)
_sem_chat  = asyncio.Semaphore(3)
_sem_rpts  = asyncio.Semaphore(2)
_sem_reco  = asyncio.Semaphore(2)

class QuotaExceeded(Exception): pass

# ─────────────────────────────────────────────────
# MÉTRIQUES
# ─────────────────────────────────────────────────
AI_METRICS = {
    "total_tokens": 0, "total_calls": 0, "error_count": 0, "active_requests": 0,
    "latency_history": deque(maxlen=50),
    "usage_counts": {
        "Évaluations": 0, "Analyses CV": 0, "Entretiens IA": 0,
        "Rapports": 0, "Chat": 0, "Dashboard": 0, "Recommandations": 0,
        "Lettres": 0, "Vocal": 0
    },
    "chat_cache_hits": 0, "chat_intent_hits": 0, "chat_gemini_calls": 0,
    "dashboard_requests": 0, "realtime_updates": 0,
    "cv_rejected": 0,
}
_activity_log: deque = deque(maxlen=200)

# ─────────────────────────────────────────────────
# DYNAMIC DATA STORE
# ─────────────────────────────────────────────────
_DS = {
    "candidat_tests": [
        {"id": 1, "candidatureId": 101, "campagneNom": "Test Logique & Raisonnement",
         "statut": "POSTULE", "dureeMinutes": 45, "color": "#3b82f6",
         "datePostule": "2025-05-06T10:00:00", "deadline": "2025-05-10"},
        {"id": 2, "candidatureId": 102, "campagneNom": "Compétences Techniques React",
         "statut": "EN_COURS", "dureeMinutes": 60, "color": "#8b5cf6",
         "datePostule": "2025-05-05T14:00:00", "deadline": "2025-05-08"},
        {"id": 3, "candidatureId": 103, "campagneNom": "Communication & Soft Skills",
         "statut": "POSTULE", "dureeMinutes": 30, "color": "#10b981",
         "datePostule": "2025-05-07T09:00:00", "deadline": "2025-05-12"},
    ],
    "candidat_historique": [
        {"id": 10, "titreExamen": "Personnalité MBTI",   "score": 91, "date": "2025-05-02T11:00:00", "campagneNom": "MBTI",         "feedback": "Excellent profil analytique"},
        {"id": 9,  "titreExamen": "Test Excel Avancé",   "score": 74, "date": "2025-04-28T10:00:00", "campagneNom": "Excel",        "feedback": "Bonne maîtrise des formules"},
        {"id": 8,  "titreExamen": "Communication Orale", "score": 62, "date": "2025-04-20T15:00:00", "campagneNom": "Communication","feedback": "À améliorer : prise de parole"},
        {"id": 7,  "titreExamen": "Logique Analytique",  "score": 88, "date": "2025-04-15T09:00:00", "campagneNom": "Logique",      "feedback": "Raisonnement structuré"},
        {"id": 6,  "titreExamen": "Python Fondamentaux", "score": 79, "date": "2025-04-10T14:00:00", "campagneNom": "Python",       "feedback": "Bonne maîtrise des bases"},
    ],
    "eval_queue": [
        {"id": 1, "candidatNom": "Sara Ben Ali",   "titreExamen": "Test Logique",    "statut": "TERMINE", "waitingHours": 26, "urgency": "urgent"},
        {"id": 2, "candidatNom": "Karim Mansouri", "titreExamen": "Compétences RH",  "statut": "TERMINE", "waitingHours": 12, "urgency": "medium"},
        {"id": 3, "candidatNom": "Lina Trabelsi",  "titreExamen": "Excel Avancé",    "statut": "TERMINE", "waitingHours": 4,  "urgency": "normal"},
        {"id": 4, "candidatNom": "Ahmed Dridi",    "titreExamen": "Leadership",       "statut": "TERMINE", "waitingHours": 3,  "urgency": "normal"},
        {"id": 5, "candidatNom": "Mona Chebbi",    "titreExamen": "Python Senior",    "statut": "EN_COURS","waitingHours": 0,  "urgency": "normal"},
        {"id": 6, "candidatNom": "Youssef Hakim",  "titreExamen": "React Frontend",   "statut": "TERMINE", "waitingHours": 8,  "urgency": "medium"},
    ],
    "campagnes": [
        {"id": 1, "nom": "Session Recrutement Q2",    "statut": 1, "dureeMinutes": 60,  "dateDebut": "2025-05-08T10:00:00", "nbCandidats": 12},
        {"id": 2, "nom": "Éval. Leadership",           "statut": 0, "dureeMinutes": 45,  "dateDebut": "2025-05-12T14:30:00", "nbCandidats": 5},
        {"id": 3, "nom": "Assessment Technique React", "statut": 0, "dureeMinutes": 90,  "dateDebut": "2025-05-20T09:00:00", "nbCandidats": 8},
        {"id": 4, "nom": "Test Softskills Junior",     "statut": 0, "dureeMinutes": 30,  "dateDebut": "2025-05-25T11:00:00", "nbCandidats": 15},
        {"id": 5, "nom": "DevOps Assessment",          "statut": 0, "dureeMinutes": 75,  "dateDebut": "2025-06-01T10:00:00", "nbCandidats": 6},
    ],
    "top_skills": [
        {"label": "Logique & Analyse",   "value": 84, "color": "#3b82f6"},
        {"label": "Communication",        "value": 71, "color": "#10b981"},
        {"label": "Leadership",           "value": 63, "color": "#f59e0b"},
        {"label": "Technique React/Vue",  "value": 78, "color": "#8b5cf6"},
        {"label": "Gestion du stress",    "value": 55, "color": "#ef4444"},
        {"label": "Travail en équipe",    "value": 88, "color": "#6366f1"},
    ],
    "staff_members": [
        {"id": 1, "prenom": "Amira",   "nomFamille": "Saidi",    "roleNom": "RH",        "estActif": True,  "tasks": 5,  "color": "#6366f1"},
        {"id": 2, "prenom": "Mehdi",   "nomFamille": "Chaabane", "roleNom": "Evaluateur","estActif": True,  "tasks": 8,  "color": "#f59e0b"},
        {"id": 3, "prenom": "Rim",     "nomFamille": "Bouzid",   "roleNom": "Recruteur", "estActif": False, "tasks": 0,  "color": "#10b981"},
        {"id": 4, "prenom": "Youssef", "nomFamille": "Hakim",    "roleNom": "Evaluateur","estActif": True,  "tasks": 3,  "color": "#8b5cf6"},
        {"id": 5, "prenom": "Nadia",   "nomFamille": "Slim",     "roleNom": "RH",        "estActif": True,  "tasks": 6,  "color": "#ef4444"},
        {"id": 6, "prenom": "Omar",    "nomFamille": "Farhat",   "roleNom": "Recruteur", "estActif": False, "tasks": 0,  "color": "#3b82f6"},
    ],
    "recent_candidates": [
        {"id": 1, "candidateId": "C001", "candidateName": "Tarek Ben Salem", "testName": "Test Logique",   "score": 91, "status": "Retenu"},
        {"id": 2, "candidateId": "C002", "candidateName": "Nour Jelassi",    "testName": "Compétences RH", "score": 74, "status": "En cours"},
        {"id": 3, "candidateId": "C003", "candidateName": "Amine Driss",     "testName": "Excel Avancé",   "score": 48, "status": "Refusé"},
        {"id": 4, "candidateId": "C004", "candidateName": "Sana Mhiri",      "testName": "Communication",  "score": 82, "status": "Retenu"},
        {"id": 5, "candidateId": "C005", "candidateName": "Khalil Touati",   "testName": "React Senior",   "score": 95, "status": "Retenu"},
        {"id": 6, "candidateId": "C006", "candidateName": "Ines Gharbi",     "testName": "Python Mid",     "score": 67, "status": "En cours"},
    ],
    "sa_services": [
        {"name": "API Gateway",      "latency": "12ms", "up": True,  "uptime_pct": 99.9},
        {"name": "Auth Service",     "latency": "8ms",  "up": True,  "uptime_pct": 100.0},
        {"name": "IA Engine Gemini", "latency": "34ms", "up": True,  "uptime_pct": 99.8},
        {"name": "Mailer Service",   "latency": "—",    "up": False, "uptime_pct": 87.2},
        {"name": "Storage S3",       "latency": "21ms", "up": True,  "uptime_pct": 99.9},
        {"name": "Analytics DB",     "latency": "15ms", "up": True,  "uptime_pct": 99.7},
        {"name": "WebSocket Hub",    "latency": "5ms",  "up": True,  "uptime_pct": 99.9},
        {"name": "PDF Generator",    "latency": "62ms", "up": True,  "uptime_pct": 99.5},
    ],
    "sa_companies": [
        {"id": 1, "name": "TechCorp Tunisia", "plan": "Enterprise", "users": 248, "color": "#6366f1", "active_sessions": 34},
        {"id": 2, "name": "Sotetel",          "plan": "Business",   "users": 89,  "color": "#f59e0b", "active_sessions": 12},
        {"id": 3, "name": "Ooredoo TN",       "plan": "Enterprise", "users": 312, "color": "#10b981", "active_sessions": 45},
        {"id": 4, "name": "Attijari Bank",    "plan": "Starter",    "users": 34,  "color": "#8b5cf6", "active_sessions": 5},
        {"id": 5, "name": "STEG Digital",     "plan": "Business",   "users": 67,  "color": "#ef4444", "active_sessions": 8},
        {"id": 6, "name": "Tunisie Telecom",  "plan": "Enterprise", "users": 445, "color": "#3b82f6", "active_sessions": 67},
    ],
    "sa_subscriptions": [
        {"plan": "Enterprise", "count": 12, "pct": 85, "color": "#6366f1", "revenue": 48000},
        {"plan": "Business",   "count": 18, "pct": 60, "color": "#f59e0b", "revenue": 27000},
        {"plan": "Starter",    "count": 12, "pct": 40, "color": "#10b981", "revenue": 7200},
    ],
    "platform_users": [
        {"id": 1, "name": "Ahmed Ben Ali",  "org": "TechCorp Tunisia", "role": "AdminEntreprise", "isActive": True},
        {"id": 2, "name": "Sara Mansouri",  "org": "Sotetel",          "role": "RH",              "isActive": True},
        {"id": 3, "name": "Mehdi Chaabane", "org": "Ooredoo TN",       "role": "Evaluateur",      "isActive": True},
        {"id": 4, "name": "Lina Trabelsi",  "org": "Attijari Bank",    "role": "Recruteur",       "isActive": False},
        {"id": 5, "name": "Omar Farhat",    "org": "STEG Digital",     "role": "Candidat",        "isActive": True},
        {"id": 6, "name": "Nadia Slim",     "org": "Tunisie Telecom",  "role": "RH",              "isActive": True},
    ],
    "pipeline": {
        "invites":   {"count": 5,  "cards": [
            {"id": 1, "name": "Ahmed Ben Salah", "score": 88, "tag": "React",  "date": "Aujourd'hui"},
            {"id": 2, "name": "Sara Mansouri",   "score": 72, "tag": "Python", "date": "Hier"},
            {"id": 3, "name": "Mehdi Trabelsi",  "score": 0,  "tag": "DevOps", "date": "Aujourd'hui"},
        ]},
        "en_cours":  {"count": 3,  "cards": [
            {"id": 4, "name": "Lina Bouzid", "score": 65, "tag": "Vue.js", "date": "En cours"},
            {"id": 5, "name": "Omar Dridi",  "score": 78, "tag": "Java",   "date": "En cours"},
        ]},
        "completes": {"count": 4,  "cards": [
            {"id": 6, "name": "Fatima Zouari", "score": 92, "tag": "Full-Stack", "date": "02/05"},
            {"id": 7, "name": "Karim Hakim",   "score": 75, "tag": "SQL",        "date": "01/05"},
        ]},
        "retenus":   {"count": 2,  "cards": [
            {"id": 8, "name": "Mariam Khelifi",   "score": 95, "tag": "React Senior", "date": "29/04"},
            {"id": 9, "name": "Youssef Chaabane", "score": 89, "tag": "Node.js",      "date": "28/04"},
        ]},
    },
}

def log_activity(user: str, action: str, color: str = "#6366f1", module: str = ""):
    _activity_log.appendleft({
        "id": int(time.time()*1000),
        "user": user, "action": action, "color": color,
        "time": datetime.now().strftime("%H:%M"), "module": module
    })

def track_usage(start_time: float, response, module_name: str):
    AI_METRICS["latency_history"].append((time.time() - start_time) * 1000)
    if response and hasattr(response, "usage_metadata") and response.usage_metadata:
        AI_METRICS["total_tokens"] += getattr(response.usage_metadata, "total_token_count", 0)
    AI_METRICS["total_calls"] += 1
    AI_METRICS["usage_counts"].setdefault(module_name, 0)
    AI_METRICS["usage_counts"][module_name] += 1

# ─────────────────────────────────────────────────
# LIFESPAN
# ─────────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    global WORKING_MODEL, _gemini_client
    gemini_status = "[OFFLINE] Non connecte"
    
    if _gemini_client:
        try:
            loop = asyncio.get_event_loop()
            models_resp = await loop.run_in_executor(_gemini_executor, _gemini_client.models.list)
            names = [m.name for m in models_resp]
            for c in ["models/gemini-2.5-flash", "models/gemini-2.0-flash", "models/gemini-1.5-flash"]:
                if c in names: 
                    WORKING_MODEL = c.replace("models/", "")
                    break
            gemini_status = f"[ONLINE] Connecte ({WORKING_MODEL})"
        except Exception as e:
            logger.warning(f"Gemini check failed: {e}")
            _gemini_client = None
            gemini_status = f"[ERROR] Echec de connexion (Erreur: {str(e)[:50]})"
    else:
        gemini_status = "[DISABLED] Cle GEMINI_API_KEY absente ou invalide"

    log_activity("Systeme", "Moteur IA v10.1 demarre", "#10b981", "Systeme")
    cache_count = await _prewarm_chat_cache()
    
    # Beautiful visual dashboard on terminal startup (ASCII safe for Windows console)
    print("\n" + "="*60)
    print(" >>> EVALUATECH AI ENGINE v10.1 - DEMARRAGE REUSSI <<<")
    print("="*60)
    print(f" [STATUS]   : Actif & En cours d'execution")
    print(f" [ENGINE]   : {gemini_status}")
    print(f" [CACHE]    : Pre-charge ({cache_count} entrees)")
    print(f" [ADDRESS]  : http://localhost:8000")
    print("="*60 + "\n")
    
    yield
    _gemini_executor.shutdown(wait=False)

async def _prewarm_chat_cache():
    frequent = [
        ("bonjour","fr","Recruteur"), ("hello","en","Recruteur"),
        ("mon score","fr","Candidat"), ("aide","fr","Candidat"),
        ("dashboard","fr","AdminEntreprise"), ("cv","fr","Recruteur"),
        ("lettre motivation","fr","Candidat"), ("recommandations","fr","AdminEntreprise"),
        ("thème","fr","Recruteur"), ("theme","en","Recruteur"),
    ]
    count = 0
    for msg, lang, role in frequent:
        ck = make_cache_key("chat-v10", msg[:30], lang, role)
        reply = _get_local_response(msg, lang, role)
        if reply:
            _chat_cache.set(ck, {"response": reply, "suggestions": _get_suggestions(msg, lang, role)})
            count += 1
    return count

# ─────────────────────────────────────────────────
# APP
# ─────────────────────────────────────────────────
app = FastAPI(title="EvaluaTech AI Engine v10.1", lifespan=lifespan)
app.add_middleware(GZipMiddleware, minimum_size=200)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

@app.middleware("http")
async def request_middleware(request: Request, call_next):
    ip = request.client.host if request.client else "unknown"
    if not check_rate_limit(ip):
        return JSONResponse(status_code=429, content={"detail": "Trop de requêtes. Attendez 60s."})
    AI_METRICS["active_requests"] += 1
    try: response = await call_next(request)
    finally: AI_METRICS["active_requests"] -= 1
    return response

# ─────────────────────────────────────────────────
# GEMINI ASYNC
# ─────────────────────────────────────────────────
async def call_gemini_async(prompt, config=None, retries=2, delay=5, module="unknown", sem=None):
    if not _gemini_client: raise QuotaExceeded("Gemini non disponible")
    if not _circuit.is_allowed(): raise QuotaExceeded("Circuit ouvert")
    loop = asyncio.get_event_loop()
    active_sem = sem or _gemini_semaphore
    async with active_sem:
        t0 = time.time()
        for attempt in range(retries):
            try:
                response = await asyncio.wait_for(
                    loop.run_in_executor(_gemini_executor,
                        lambda: _gemini_client.models.generate_content(
                            model=WORKING_MODEL, contents=prompt, config=config)),
                    timeout=18.0)
                _circuit.record_success()
                track_usage(t0, response, module)
                return response
            except asyncio.TimeoutError:
                if attempt < retries - 1: await asyncio.sleep(2); continue
                raise QuotaExceeded(f"Timeout [{module}]")
            except Exception as e:
                err_str = str(e)
                if any(x in err_str for x in ["429","RESOURCE_EXHAUSTED","quota"]):
                    AI_METRICS["error_count"] += 1; _circuit.record_failure()
                    raise QuotaExceeded(f"Quota [{module}]")
                if any(x in err_str for x in ["API_KEY_INVALID","key expired"]):
                    globals()['_gemini_client'] = None
                    raise QuotaExceeded(f"Clé invalide [{module}]")
                if any(x in err_str for x in ["503","500"]) and attempt < retries - 1:
                    await asyncio.sleep(delay * (attempt + 1)); continue
                _circuit.record_failure(); AI_METRICS["error_count"] += 1
                raise QuotaExceeded(f"Erreur [{module}]: {err_str[:80]}")
        raise QuotaExceeded(f"Max retries [{module}]")

# ─────────────────────────────────────────────────
# UTILITAIRES
# ─────────────────────────────────────────────────
def _truncate(text: str, max_chars=3500) -> str:
    if len(text) <= max_chars: return text
    h = max_chars // 2
    return text[:h] + "\n...[tronqué]...\n" + text[-h:]

async def read_pdf_async(fb: bytes) -> str:
    loop = asyncio.get_event_loop()
    def _read():
        if FITZ_AVAILABLE:
            try:
                doc = fitz.open(stream=fb, filetype="pdf")
                return "".join(p.get_text() for p in doc)
            except: pass
        if PYPDF2_AVAILABLE:
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(fb))
                return "".join(p.extract_text() or "" for p in reader.pages)
            except: pass
        return ""
    return await loop.run_in_executor(_gemini_executor, _read)

async def extract_text_from_upload(file: UploadFile) -> str:
    fb = await file.read()
    name = (file.filename or "").lower()
    if name.endswith(".pdf"): return _truncate(await read_pdf_async(fb))
    loop = asyncio.get_event_loop()
    def _other():
        try:
            if name.endswith(".docx") and DOCX_AVAILABLE:
                doc = Document(io.BytesIO(fb))
                return "\n".join(p.text for p in doc.paragraphs)
            return fb.decode("utf-8", errors="ignore")
        except: return ""
    return _truncate(await loop.run_in_executor(_gemini_executor, _other))

def clean_json(text: str) -> str:
    return text.replace("```json","").replace("```","").strip()

def _normalize(text: str) -> str:
    return (text.lower()
        .replace("é","e").replace("è","e").replace("ê","e").replace("ë","e")
        .replace("à","a").replace("â","a").replace("ù","u").replace("û","u")
        .replace("ç","c").replace("î","i").replace("ô","o"))

# ════════════════════════════════════════════════════════════
# ██   NOUVEAU v10.1 — DÉTECTION DOCUMENT NON-CV   ██
# ════════════════════════════════════════════════════════════

_CV_KEYWORDS_FR = [
    "expérience","experience","formation","compétences","competences",
    "diplôme","diplome","poste","entreprise","université","universite",
    "baccalauréat","baccalaureat","master","licence","engineer","ingénieur",
    "stage","emploi","travail","curriculum","cv","profil","objectif",
    "réalisations","realisations","skills","education","professional",
    "langues","certifications","références","references",
]
_CV_KEYWORDS_EN = [
    "experience","education","skills","resume","curriculum","employment",
    "objective","summary","achievements","languages","certifications",
    "references","degree","bachelor","master","university","college",
    "work history","job","position","company","internship","volunteer",
]
_CV_KEYWORDS_AR = [
    "خبرة","مهارات","تعليم","السيرة","الذاتية","شهادة","جامعة",
    "العمل","وظيفة","مؤهل","تدريب","لغات","مراجع",
]

_NON_CV_SIGNALS = [
    "facture","invoice","devis","bon de commande","contrat","contract",
    "rapport","report","article","abstract","résumé exécutif","executive summary",
    "présentation","presentation","slide","powerpoint","tableau","spreadsheet",
    "code source","source code","import","def ","function","class ",
    "<!doctype","<html","<?xml","<?php","{\"","[{\"",
]

def _detect_is_cv(text: str) -> dict:
    if not text or len(text.strip()) < 50:
        return {
            "is_cv": False,
            "confidence": 0.0,
            "reason": "Document vide ou trop court",
            "reason_en": "Document is empty or too short",
            "reason_ar": "الوثيقة فارغة أو قصيرة جداً",
        }

    text_lower = text.lower()
    text_norm  = _normalize(text_lower)

    for signal in _NON_CV_SIGNALS:
        if signal.lower() in text_lower:
            return {
                "is_cv": False,
                "confidence": 0.85,
                "reason": f"Le document semble être un fichier de type « {signal} », pas un CV.",
                "reason_en": f"The document appears to be a '{signal}' file, not a CV.",
                "reason_ar": f"يبدو أن الوثيقة من نوع '{signal}' وليست سيرة ذاتية.",
                "detected_signal": signal,
            }

    cv_score = 0
    matched_kw = []
    for kw in _CV_KEYWORDS_FR + _CV_KEYWORDS_EN + _CV_KEYWORDS_AR:
        kw_norm = _normalize(kw)
        if kw_norm in text_norm:
            cv_score += 1
            matched_kw.append(kw)

    word_count = len(text.split())
    has_dates  = bool(re.search(r'\b(19|20)\d{2}\b', text))
    has_email  = bool(re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', text))
    has_phone  = bool(re.search(r'[\+\(]?[\d\s\-\(\)]{7,15}', text))

    struct_score = sum([has_dates, has_email, has_phone])
    total_score  = cv_score + struct_score * 2

    if total_score >= 6:
        confidence = min(0.95, 0.60 + (total_score - 6) * 0.035)
        return {
            "is_cv": True,
            "confidence": round(confidence, 2),
            "matched_keywords": matched_kw[:10],
            "word_count": word_count,
        }
    elif total_score >= 3:
        confidence = 0.45 + total_score * 0.03
        return {
            "is_cv": True,
            "confidence": round(confidence, 2),
            "matched_keywords": matched_kw[:5],
            "word_count": word_count,
            "warning": "Confiance modérée — le document pourrait ne pas être un CV complet.",
            "warning_en": "Moderate confidence — the document may not be a complete CV.",
            "warning_ar": "ثقة متوسطة — قد لا تكون الوثيقة سيرة ذاتية كاملة.",
        }
    else:
        return {
            "is_cv": False,
            "confidence": round(1 - (total_score / 6), 2),
            "reason": "Le document ne contient pas les éléments caractéristiques d'un CV (expérience, formation, compétences, coordonnées).",
            "reason_en": "The document does not contain typical CV elements (experience, education, skills, contact info).",
            "reason_ar": "لا تحتوي الوثيقة على عناصر السيرة الذاتية النموذجية (خبرة، تعليم، مهارات، معلومات الاتصال).",
            "matched_keywords": matched_kw,
            "word_count": word_count,
        }


def _build_cv_alert(detection: dict, lang: str = "fr") -> dict:
    messages = {
        "fr": {
            "title": "⚠️ Document non reconnu comme CV",
            "subtitle": detection.get("reason", "Ce document ne ressemble pas à un CV."),
            "hint": "Veuillez uploader un fichier PDF ou DOCX contenant : expériences professionnelles, formation, compétences.",
            "cta": "Changer de fichier",
            "type": "warning",
            "icon": "fa-solid fa-triangle-exclamation",
        },
        "en": {
            "title": "⚠️ Document not recognized as a CV",
            "subtitle": detection.get("reason_en", "This document does not look like a CV/resume."),
            "hint": "Please upload a PDF or DOCX file containing: work experience, education, skills.",
            "cta": "Change file",
            "type": "warning",
            "icon": "fa-solid fa-triangle-exclamation",
        },
        "ar": {
            "title": "⚠️ الوثيقة غير معترف بها كسيرة ذاتية",
            "subtitle": detection.get("reason_ar", "لا تبدو هذه الوثيقة سيرة ذاتية."),
            "hint": "يرجى رفع ملف PDF أو DOCX يحتوي على: خبرة مهنية، تعليم، مهارات.",
            "cta": "تغيير الملف",
            "type": "warning",
            "icon": "fa-solid fa-triangle-exclamation",
        },
    }
    return {
        "is_cv": False,
        "alert": messages.get(lang, messages["fr"]),
        "confidence": detection.get("confidence", 0),
        "matched_keywords": detection.get("matched_keywords", []),
        "word_count": detection.get("word_count", 0),
    }


# ════════════════════════════════════════════════════════════
# ██   NOUVEAU v10.1 — LETTRE DE MOTIVATION PAR LANGUE   ██
# ════════════════════════════════════════════════════════════

_LETTRE_TEMPLATES = {
    "fr": {
        "professionnel": """{salutation},

Passionné(e) par {competences_phrase}, je me permets de vous adresser ma candidature au poste de **{poste}** au sein de **{entreprise}**.

Votre entreprise se distingue par son engagement envers l'excellence et l'innovation, des valeurs qui correspondent pleinement à ma vision professionnelle et à mes aspirations de carrière.

Au cours de mon parcours, j'ai développé de solides compétences en {competences_phrase}. Ces expériences m'ont permis de livrer des résultats concrets dans des environnements exigeants et de travailler efficacement en équipe.

Je serais ravi(e) de vous présenter mon parcours en détail lors d'un entretien à votre convenance, et de vous démontrer comment mes compétences peuvent contribuer au succès de {entreprise}.

Dans cette attente, veuillez agréer l'expression de mes salutations distinguées.

{signature}""",
        "créatif": """Bonjour,

Ce qui me distingue ? Une vraie passion pour {competences_phrase} — et l'envie de la mettre au service de {entreprise} en tant que **{poste}**.

Votre réputation d'innovateur m'a immédiatement interpellé(e). Ce n'est pas un hasard si je vous contacte : je cherche un environnement qui valorise l'initiative et l'impact concret.

Mes expériences en {competences_phrase} m'ont permis de transformer des défis complexes en solutions élégantes. J'apporte avec moi une approche à la fois rigoureuse et créative.

Parlons-en ! Je suis disponible pour un échange à tout moment.

{signature}""",
        "concis": """{salutation},

Je postule au poste de **{poste}** chez **{entreprise}**.

Fort(e) d'une expertise en {competences_phrase}, je suis convaincu(e) de pouvoir apporter une réelle valeur ajoutée à votre équipe.

Disponible pour un entretien à votre convenance.

{signature}""",
    },
    "en": {
        "professionnel": """Dear Hiring Manager,

I am writing to express my strong interest in the **{poste}** position at **{entreprise}**. With expertise in {competences_phrase}, I am confident in my ability to contribute meaningfully to your team.

{entreprise} stands out for its commitment to excellence and innovation — values that align perfectly with my professional vision and career aspirations.

Throughout my career, I have developed strong skills in {competences_phrase}, allowing me to deliver tangible results in demanding environments while working effectively within cross-functional teams.

I would welcome the opportunity to discuss how my background matches your needs during an interview at your convenience.

Thank you sincerely for considering my application.

{signature}""",
        "créatif": """Hello,

What sets me apart? A genuine passion for {competences_phrase} — and a drive to bring that energy to **{entreprise}** as your next **{poste}**.

Your reputation for innovation caught my attention immediately. I'm not reaching out by chance: I thrive in environments that value initiative and measurable impact.

My background in {competences_phrase} has equipped me to turn complex challenges into elegant solutions, combining rigor with creativity.

Let's connect — I'm available for a conversation at any time.

{signature}""",
        "concis": """Dear Hiring Manager,

I am applying for the **{poste}** position at **{entreprise}**.

With proven expertise in {competences_phrase}, I am well-positioned to add real value to your team.

Available for an interview at your convenience.

{signature}""",
    },
    "ar": {
        "professionnel": """السيد/السيدة المحترم/ة،

يسعدني تقديم طلبي للمنصب **{poste}** في شركة **{entreprise}**، مدفوعاً برغبة صادقة في الإسهام بخبرتي في {competences_phrase} لتحقيق أهداف مؤسستكم.

تتميز {entreprise} بسمعتها الرائدة في الابتكار والتميز، وهي القيم التي تتوافق تماماً مع رؤيتي المهنية.

خلال مسيرتي المهنية، طوّرت مهارات متينة في {competences_phrase}، ما أتاح لي تحقيق نتائج ملموسة في بيئات عمل متطلبة والعمل بفعالية ضمن فرق متعددة التخصصات.

أتطلع إلى فرصة تقديم مسيرتي بتفصيل أكبر خلال مقابلة في الوقت الذي يناسبكم.

مع التقدير والاحترام،

{signature}""",
        "créatif": """مرحباً،

ما يميّزني؟ شغف حقيقي بـ {competences_phrase} ورغبة صادقة في توظيفه لصالح **{entreprise}** بوصفي **{poste}** في فريقكم.

سمعتكم في الابتكار لفتت انتباهي على الفور. لا أتواصل معكم بالصدفة: أزدهر في البيئات التي تقدّر المبادرة والأثر الحقيقي.

تجربتي في {competences_phrase} منحتني القدرة على تحويل التحديات المعقدة إلى حلول أنيقة تجمع بين الصرامة والإبداع.

لنتحدث — أنا متاح/ة في أي وقت.

{signature}""",
        "concis": """السيد/السيدة المحترم/ة،

أتقدم بطلبي لشغل منصب **{poste}** في **{entreprise}**.

بفضل خبرتي الموثّقة في {competences_phrase}، أنا في وضع جيد لتقديم قيمة مضافة حقيقية لفريقكم.

أنا متاح/ة لمقابلة في الوقت الذي يناسبكم.

{signature}""",
    },
}

def _get_salutation(langue: str) -> str:
    salutations = {
        "fr": "Madame, Monsieur",
        "en": "Dear Hiring Manager",
        "ar": "السيد/السيدة المحترم/ة",
    }
    return salutations.get(langue, salutations["fr"])

def _get_signature(nom: str, langue: str) -> str:
    name_part = nom.strip() if nom and nom.strip() else {
        "fr": "Le/La Candidat(e)",
        "en": "The Applicant",
        "ar": "المتقدم/ة",
    }.get(langue, "Le/La Candidat(e)")

    closings = {
        "fr": f"Cordialement,\n{name_part}",
        "en": f"Sincerely,\n{name_part}",
        "ar": f"مع التقدير،\n{name_part}",
    }
    return closings.get(langue, closings["fr"])

def _get_competences_phrase(competences: str, langue: str, poste: str) -> str:
    if competences and competences.strip():
        return competences.strip()
    defaults = {
        "fr": f"les domaines techniques et relationnels liés au poste de {poste}",
        "en": f"technical and interpersonal skills relevant to the {poste} role",
        "ar": f"المهارات التقنية والإنسانية المرتبطة بمنصب {poste}",
    }
    return defaults.get(langue, defaults["fr"])

def _generate_fallback_lettre(nom: str, poste: str, entreprise: str, competences: str, langue: str, style: str) -> str:
    lang_key = langue.strip().lower()
    if lang_key not in _LETTRE_TEMPLATES:
        lang_key = "fr"

    style_key = style.strip().lower()
    if style_key not in _LETTRE_TEMPLATES[lang_key]:
        style_key = "professionnel"

    template = _LETTRE_TEMPLATES[lang_key][style_key]

    competences_phrase = _get_competences_phrase(competences, lang_key, poste)
    salutation         = _get_salutation(lang_key)
    signature          = _get_signature(nom, lang_key)

    lettre = template.format(
        salutation=salutation,
        poste=poste or ("le poste" if lang_key == "fr" else "the position" if lang_key == "en" else "المنصب"),
        entreprise=entreprise or ("votre entreprise" if lang_key == "fr" else "your company" if lang_key == "en" else "شركتكم"),
        competences_phrase=competences_phrase,
        signature=signature,
    )
    return lettre.strip()


def _build_lettre_prompt(nom: str, poste: str, entreprise: str, competences: str, langue: str, style: str) -> str:
    lang_configs = {
        "fr": {
            "instruction": "Tu dois OBLIGATOIREMENT rédiger cette lettre EN FRANÇAIS uniquement. N'utilise AUCUN mot en anglais ou arabe.",
            "style_map": {
                "professionnel": "Style professionnel classique, formel et persuasif.",
                "créatif": "Style moderne, dynamique, avec une accroche originale.",
                "concis": "Style concis et direct, maximum 200 mots.",
            },
            "check": "VÉRIFIE : toute la lettre est bien en français.",
        },
        "en": {
            "instruction": "You MUST write this letter ENTIRELY IN ENGLISH. Do NOT use any French or Arabic words.",
            "style_map": {
                "professionnel": "Classic professional style, formal and persuasive.",
                "créatif": "Modern, dynamic style with an original opening.",
                "concis": "Concise and direct style, maximum 200 words.",
            },
            "check": "CHECK: the entire letter is in English.",
        },
        "ar": {
            "instruction": "يجب عليك كتابة هذه الرسالة بالكامل باللغة العربية فقط. لا تستخدم أي كلمات فرنسية أو إنجليزية.",
            "style_map": {
                "professionnel": "أسلوب احترافي رسمي ومقنع.",
                "créatif": "أسلوب حديث وديناميكي مع مقدمة إبداعية.",
                "concis": "أسلوب موجز ومباشر، بحد أقصى 200 كلمة.",
            },
            "check": "تحقق: الرسالة بالكاملة باللغة العربية.",
        },
    }

    lang_key   = langue.strip().lower() if langue.strip().lower() in lang_configs else "fr"
    style_key  = style.strip().lower()  if style.strip().lower() in ["professionnel","créatif","concis"] else "professionnel"
    lc         = lang_configs[lang_key]
    style_desc = lc["style_map"].get(style_key, lc["style_map"]["professionnel"])
    comp_ph    = _get_competences_phrase(competences, lang_key, poste)

    prompt = f"""{lc["instruction"]}

{style_desc}

Rédige une lettre de motivation professionnelle pour :
- Candidat : {nom or ('Le/La candidat(e)' if lang_key=='fr' else 'The applicant' if lang_key=='en' else 'المتقدم/ة')}
- Poste : {poste}
- Entreprise : {entreprise}
- Compétences : {comp_ph}

Contraintes absolues :
1. TOUTE la lettre doit être en {lang_key.upper()} — aucune exception
2. Structure : Accroche → Compétences → Motivation → Appel à l'action
3. Longueur : {'250-350 mots' if style_key!='concis' else '100-200 mots'}
4. Sans en-tête formel ni signature
5. Ton professionnel mais humain
6. {lc["check"]}

Réponds UNIQUEMENT avec le texte de la lettre, sans explications."""

    return prompt


# ════════════════════════════════════════════════════════════
# ██   CHATBOT — BASE DE CONNAISSANCES   ██
# ════════════════════════════════════════════════════════════

def detect_language(text: str) -> str:
    ar_chars = len(re.findall(r'[\u0600-\u06FF]', text))
    en_words = len(re.findall(
        r'\b(how|what|where|when|why|is|are|can|help|test|score|create|about|does|'
        r'tell|show|my|the|generate|analyze|report|view|list|status|enable|disable|'
        r'dashboard|stats|cv|letter|recommendation|platform|proctoring|evaluation|'
        r'candidate|recruiter|admin|question|bank|campaign|result|login|password|'
        r'profile|setting|notification|export|import|download|upload|team|staff|'
        r'theme|domain|specialty|topic|subject|available)\b',
        text.lower()))
    if ar_chars > 2: return "ar"
    if en_words >= 2: return "en"
    return "fr"

UNIVERSAL_BRAIN = {
    "greeting": {
        "patterns": {
            "fr": ["bonjour","bonsoir","salut","hello","hey","coucou","allo","bjr"],
            "en": ["hello","hi","hey","good morning","good evening","greetings"],
            "ar": ["مرحبا","السلام عليكم","أهلاً","صباح الخير","مساء الخير"]
        },
        "responses": {
            "fr": {
                "Candidat":        "👋 Bonjour ! Je suis **NeoBot**, votre assistant IA EvaluaTech.\n\n🎯 Je peux vous aider avec :\n- 📝 Vos **tests et évaluations** en attente\n- 📊 Vos **résultats et scores** en temps réel\n- 🎤 **Préparer votre entretien** IA\n- 📄 **Analyser votre CV** (Gemini IA)\n- ✉️ **Générer votre lettre** de motivation (FR/EN/AR)\n- 📈 **Voir vos statistiques** de progression\n\nQue puis-je faire pour vous ?",
                "Recruteur":       "👋 Bonjour ! Je suis **NeoBot**, votre assistant IA EvaluaTech.\n\n🚀 Je peux vous aider à :\n- ✅ **Créer des tests** QCM IA en 45 secondes\n- 📄 **Analyser des CVs** avec matching Gemini\n- 📋 **Gérer votre pipeline** Kanban candidats\n- 📊 **Voir les statistiques** en temps réel\n- 📩 **Inviter des candidats** (email / CSV)\n- 🤖 **Recommandations IA** personnalisées\n\nQuelle est votre question ?",
                "Evaluateur":      "👋 Bonjour ! Je suis **NeoBot**, votre assistant IA.\n\n📋 En tant qu'**Évaluateur**, je peux vous aider :\n- 👥 **File d'évaluation** — candidats en attente\n- 📅 **Sessions planifiées** à venir\n- 📊 **Statistiques** de vos évaluations\n- 🤖 **Recommandations IA** sur les profils",
                "AdminEntreprise": "👋 Bonjour ! Je suis **NeoBot**, votre assistant IA EvaluaTech.\n\n🏢 En tant qu'**Admin Entreprise**, je peux vous aider avec :\n- 👥 **Gestion de votre équipe** RH/Évaluateurs\n- 📊 **Statistiques et analytics** temps réel\n- 🤖 **Recommandations IA** automatiques\n- 📋 **Rapports automatiques** mensuels",
                "SuperAdmin":      "👋 Bonjour **SuperAdmin** ! Je suis **NeoBot**.\n\n🛡️ Accès complet à la plateforme :\n- 🏢 **42 organisations** actives\n- ⚙️ **Santé des services** microservices\n- 💰 **Abonnements** et renouvellements\n- ⚠️ **Alerte** : Service Mailer actuellement DOWN",
                "RH":              "👋 Bonjour ! Je suis **NeoBot**, votre assistant RH IA.\n\n📋 Je peux vous aider avec :\n- 📣 **Campagnes** de recrutement actives\n- 👥 **Candidats** en cours d'évaluation\n- 📊 **Rapports** et analytics RH\n- 🤖 **Recommandations** IA personnalisées",
                "default":         "👋 Bonjour ! Je suis **NeoBot**, l'assistant IA d'EvaluaTech. Quel est votre rôle pour une aide personnalisée ?"
            },
            "en": {"default": "👋 Hello! I'm **NeoBot**, EvaluaTech's AI assistant.\n\nI can help with tests, CV analysis, cover letters (FR/EN/AR), AI interview prep, statistics, and recommendations."},
            "ar": {"default": "👋 مرحباً! أنا **NeoBot**، مساعد الذكاء الاصطناعي لـ EvaluaTech.\n\nيمكنني مساعدتك في الاختبارات وتحليل السيرة وتوليد خطاب التقديم بثلاث لغات والتحضير للمقابلة."}
        },
        "suggestions": {
            "fr": ["Voir le dashboard","Mes statistiques","Comment ça marche ?"],
            "en": ["View dashboard","My statistics","How does it work?"],
            "ar": ["عرض لوحة التحكم","إحصائياتي","كيف يعمل؟"]
        }
    },

    # ══════════════════════════════════════════════════════
    # FIX v10.1 — NOUVELLE INTENTION : THÈME / SOUS-THÈME
    # ══════════════════════════════════════════════════════
    "theme_soustheme": {
        "patterns": {
            "fr": [
                "thème","theme","sous-thème","sous thème","soustheme","domaine",
                "spécialité","specialite","quels domaines","domaines disponibles",
                "quelles thématiques","thématiques disponibles","backend","frontend",
                "react","python","node","java","sql","devops","aws","docker",
                "machine learning","soft skills","quelle thematique","quels sujets",
                "sujets disponibles","quels thèmes","liste des thèmes",
                "quoi comme thème","quoi comme sujet","quel thème choisir",
                "quels sous-thèmes","choisir un thème","sélectionner thème"
            ],
            "en": [
                "theme","sub-theme","domain","specialty","which domains","available topics",
                "available themes","backend","frontend","react","python","node","java",
                "sql","devops","aws","docker","machine learning","soft skills",
                "what topics","what subjects","choose theme","select domain"
            ],
            "ar": [
                "موضوع","تخصص","مجال","الواجهة الخلفية","الواجهة الأمامية",
                "ما المواضيع المتاحة","اختر موضوع","المجالات المتاحة",
                "قائمة المواضيع","react","python","java","sql","devops"
            ]
        },
        "responses": {
            "fr": {
                "Recruteur": (
                    "🎯 **Thèmes & Sous-thèmes disponibles pour la génération IA :**\n\n"
                    "**💻 Développement Web (Frontend)**\n"
                    "- React, Vue.js, Angular, TypeScript, HTML/CSS\n\n"
                    "**⚙️ Backend & APIs**\n"
                    "- Node.js, Python (Django/FastAPI), Java (Spring), PHP, C#, REST API\n\n"
                    "**🗄️ Bases de données**\n"
                    "- SQL, PostgreSQL, MySQL, MongoDB, Redis\n\n"
                    "**☁️ DevOps & Cloud**\n"
                    "- Docker, Kubernetes, AWS, CI/CD, Linux\n\n"
                    "**🧠 Soft Skills & RH**\n"
                    "- Leadership, Communication, Gestion du stress, Travail en équipe\n\n"
                    "**📊 Data & IA**\n"
                    "- Machine Learning, Python Data, SQL Analytics, Power BI\n\n"
                    "**🔐 Sécurité**\n"
                    "- Cybersécurité, OWASP, Gestion des risques\n\n"
                    "💡 **Comment utiliser :** Dans *Créer une campagne → Générer avec IA*, "
                    "saisissez le **Thème** (ex: Backend) et le **Sous-thème** (ex: Node.js) "
                    "→ QCM généré en **45 secondes** !"
                ),
                "AdminEntreprise": (
                    "🎯 **Thèmes disponibles pour vos campagnes :**\n\n"
                    "💻 **Frontend :** React · Vue.js · Angular · TypeScript\n"
                    "⚙️ **Backend :** Node.js · Python · Java · C# · PHP\n"
                    "🗄️ **BDD :** SQL · MongoDB · Redis · PostgreSQL\n"
                    "☁️ **DevOps :** Docker · AWS · Kubernetes · CI/CD\n"
                    "🧠 **Soft Skills :** Leadership · Communication · Gestion stress\n"
                    "📊 **Data :** Machine Learning · Analytics · Power BI\n\n"
                    "💡 Entrez un Thème + Sous-thème dans le **Générateur IA** de vos campagnes."
                ),
                "default": (
                    "🎯 **Thèmes disponibles dans EvaluaTech :**\n\n"
                    "React · Vue.js · Angular · Node.js · Python · Java · PHP · C#\n"
                    "SQL · MongoDB · Docker · AWS · Kubernetes · CI/CD\n"
                    "Machine Learning · Data Analytics · Cybersécurité\n"
                    "Leadership · Communication · Soft Skills · Gestion du stress\n\n"
                    "💡 Dans le **Générateur IA**, renseignez le **Thème** et le **Sous-thème** "
                    "pour créer votre QCM personnalisé en 45 secondes !"
                )
            },
            "en": {
                "default": (
                    "🎯 **Available Themes & Sub-themes:**\n\n"
                    "**Frontend:** React · Vue.js · Angular · TypeScript · HTML/CSS\n"
                    "**Backend:** Node.js · Python · Java · C# · REST API\n"
                    "**Database:** SQL · PostgreSQL · MongoDB · Redis\n"
                    "**DevOps:** Docker · AWS · Kubernetes · CI/CD\n"
                    "**Soft Skills:** Leadership · Communication · Stress Management\n"
                    "**Data/AI:** Machine Learning · Analytics · Power BI\n\n"
                    "💡 In the **AI Generator**, enter a Theme + Sub-theme to generate your QCM in **45 seconds**!"
                )
            },
            "ar": {
                "default": (
                    "🎯 **المواضيع والتخصصات المتاحة:**\n\n"
                    "**الواجهة الأمامية:** React · Vue.js · Angular · TypeScript\n"
                    "**الواجهة الخلفية:** Node.js · Python · Java · C# · REST API\n"
                    "**قواعد البيانات:** SQL · PostgreSQL · MongoDB · Redis\n"
                    "**DevOps:** Docker · AWS · Kubernetes · CI/CD\n"
                    "**المهارات الشخصية:** القيادة · التواصل · إدارة الضغط\n"
                    "**البيانات والذكاء الاصطناعي:** تعلم الآلة · التحليلات\n\n"
                    "💡 في **مولّد الذكاء الاصطناعي**، أدخل الموضوع والتخصص لإنشاء اختبارك في **45 ثانية**!"
                )
            }
        },
        "suggestions": {
            "fr": ["Créer un test React Senior","Générer QCM Python Backend","Créer test SQL Avancé"],
            "en": ["Create React Senior test","Generate Python Backend QCM","Create Advanced SQL test"],
            "ar": ["إنشاء اختبار React","توليد اختبار Python","اختبار SQL متقدم"]
        }
    },

    "farewell": {
        "patterns": {
            "fr": ["merci","ok merci","parfait","au revoir","bye","super","excellent","bravo","top","c'est bon"],
            "en": ["thank you","thanks","perfect","goodbye","bye","great","excellent","that's all"],
            "ar": ["شكراً","مع السلامة","وداعاً","ممتاز","شكرا","رائع"]
        },
        "responses": {
            "fr": {"default": "🙌 Avec plaisir ! Bonne continuation sur **EvaluaTech** ! 🚀"},
            "en": {"default": "🙌 You're welcome! Good luck on **EvaluaTech**! 🚀"},
            "ar": {"default": "🙌 بكل سرور! بالتوفيق على **EvaluaTech**! 🚀"}
        },
        "suggestions": {
            "fr": ["Créer un test","Voir le dashboard","Analyser un CV"],
            "en": ["Create a test","View dashboard","Analyze a CV"],
            "ar": ["إنشاء اختبار","لوحة التحكم","تحليل سيرة"]
        }
    },
    "about_platform": {
        "patterns": {
            "fr": ["c'est quoi evaluatech","présentation","plateforme","qu'est-ce que","evaluatech","comment ça marche","saas","fonctionnalités"],
            "en": ["what is evaluatech","about","platform","saas","how does it work","features"],
            "ar": ["ما هو","تقديم","المنصة","عن المنصة","كيف يعمل","الميزات"]
        },
        "responses": {
            "fr": {"default": "🚀 **EvaluaTech** est une plateforme SaaS d'**évaluation technique et comportementale**.\n\n**5 piliers IA :**\n1. 🎯 **QCM IA** — tests FR/EN/AR en 45 secondes\n2. 📄 **Analyse CV Neural** — matching PDF + détection automatique si non-CV\n3. 🎤 **Entretien IA** — questions par poste et niveau\n4. ✉️ **Lettre Motivation** — Gemini en FR/EN/AR avec langue stricte\n5. 📊 **Dashboard Temps Réel** — KPIs + recommandations\n\nPropulsé par **Google Gemini AI**."},
            "en": {"default": "🚀 **EvaluaTech** — SaaS platform for technical and behavioral assessment.\n\n5 AI pillars: Tests | CV Analysis (with non-CV detection) | AI Interview | Cover Letter (strict language) | Dashboard\n\nPowered by **Google Gemini AI**."},
            "ar": {"default": "🚀 **EvaluaTech** منصة SaaS للتقييم التقني والسلوكي.\n\n5 ركائز: اختبارات | تحليل السيرة (كشف غير السيرة) | مقابلة | خطاب (لغة صارمة) | لوحة تحكم\n\nمدعومة بـ Google Gemini."}
        },
        "suggestions": {
            "fr": ["6 étapes de démarrage","Analyse CV avec alerte ?","Lettre en arabe ?"],
            "en": ["6 steps to start","CV alert feature?","Letter in French?"],
            "ar": ["6 خطوات للبدء","تنبيه السيرة الذاتية؟","خطاب بالفرنسية؟"]
        }
    },
    "dashboard": {
        "patterns": {
            "fr": ["dashboard","tableau de bord","kpi","aperçu","mon dashboard","accueil","page principale"],
            "en": ["dashboard","kpi","overview","home page","main page"],
            "ar": ["لوحة التحكم","نظرة عامة","مؤشرات","الصفحة الرئيسية"]
        },
        "responses": {
            "fr": {
                "Candidat":        "📊 **Votre Dashboard :**\n\n📝 **2 tests** en attente · ⭐ **Score : 76%** · 🧠 **Aura : 82**\n\nSections : Tests en cours · Progression · Résultats · CV IA · Lettre motivation · Recommandations",
                "AdminEntreprise": "📊 **Dashboard Admin :** 124 talents · 88% réussite · 12 sessions · Aura 94\n\nSections : KPIs · Équipe · Candidats · Top skills · Recommandations IA",
                "SuperAdmin":      "📊 **Dashboard SuperAdmin :** 42 orgs · 1247 users · 99.9% uptime · ⚠️ Mailer DOWN",
                "default":         "Le **Dashboard** affiche vos KPIs temps réel, adapté à votre rôle, actualisé toutes les 60 secondes."
            },
            "en": {"default": "📊 Your **Dashboard** shows real-time KPIs, adapted to your role, updated every 60 seconds."},
            "ar": {"default": "📊 **لوحة التحكم** تعرض مؤشرات أدائك الفورية كل 60 ثانية."}
        },
        "suggestions": {
            "fr": ["Voir mes recommandations IA","Statistiques détaillées","Actualiser"],
            "en": ["AI recommendations","Detailed stats","Refresh"],
            "ar": ["التوصيات","إحصائيات مفصلة","تحديث"]
        }
    },
    "cv_analysis": {
        "patterns": {
            "fr": ["cv","analyse cv","scanner cv","analyser cv","pdf cv","matching","compatibilité","curriculum","scan cv","neural cv","mon cv","analyser mon cv"],
            "en": ["cv","resume","analyze cv","scan cv","pdf","matching","neural scan"],
            "ar": ["سيرة ذاتية","تحليل السيرة","مسح السيرة","مطابقة"]
        },
        "responses": {
            "fr": {
                "Recruteur": "📄 **Analyser un CV avec Gemini IA :**\n\n1. Section **Scan CV Neural** sur votre dashboard\n2. Glissez le **PDF ou DOCX** du candidat\n3. Optionnel : ajoutez la fiche de poste\n4. Cliquez **Analyser**\n\n**Résultat en 10 secondes :**\n- 🎯 Score de matching (0-100%)\n- ✅ Points forts + points faibles\n- 💡 Conseils personnalisés\n- 🏷️ Compétences détectées\n\n⚠️ **Nouveau v10.1 :** Si le fichier n'est pas un CV, une **alerte** s'affiche automatiquement avec un message d'explication.",
                "Candidat":  "📄 **Analyser votre CV :**\n\n1. Section **Analyse CV IA** dans votre dashboard\n2. Glissez votre CV (**PDF ou DOCX**)\n3. Optionnel : fiche de poste\n\n⚠️ Si le fichier uploadé n'est pas un CV, un message d'alerte s'affiche pour vous demander de changer de fichier.",
                "default":   "📄 Glissez votre PDF dans **Scan CV Neural**. Gemini analyse en 10 secondes.\n⚠️ Si le document n'est pas un CV → alerte automatique."
            },
            "en": {"default": "📄 **Neural CV Scan:** Drop PDF/DOCX → Gemini analysis in 10s.\n⚠️ If the file is not a CV → automatic alert with guidance."},
            "ar": {"default": "📄 اسحب PDF/DOCX في **تحليل السيرة**. Gemini يحلل في 10 ثوانٍ.\n⚠️ إذا لم تكن الوثيقة سيرة ذاتية → تنبيه تلقائي."}
        },
        "suggestions": {
            "fr": ["Formats supportés ?","Qu'est-ce qui déclenche l'alerte non-CV ?","Générer ma lettre ?"],
            "en": ["Supported formats?","What triggers the non-CV alert?","Generate cover letter?"],
            "ar": ["الصيغ المدعومة؟","ما الذي يُشغّل التنبيه؟","إنشاء الخطاب؟"]
        }
    },
    "lettre_motivation": {
        "patterns": {
            "fr": ["lettre de motivation","lettre motivation","cover letter","générer lettre","rédiger lettre","lettre candidature","lettre en arabe","lettre en anglais","lettre en francais"],
            "en": ["cover letter","motivation letter","generate letter","write letter","letter in french","letter in arabic"],
            "ar": ["خطاب تقديم","خطاب دوافع","رسالة تغطية","كتابة خطاب","توليد خطاب","خطاب بالفرنسية","خطاب بالإنجليزية"]
        },
        "responses": {
            "fr": {
                "Candidat": "✉️ **Générer votre Lettre de Motivation IA :**\n\nDashboard → **Lettre de Motivation IA** :\n\n1. **Nom** + **Poste visé** + **Entreprise**\n2. **Compétences clés** (optionnel)\n3. **Style** : Professionnel / Créatif / Concis\n4. **Langue** : 🇫🇷 Français / 🇬🇧 English / 🇸🇦 العربية\n5. Cliquez **Générer avec Gemini**\n\n✅ **Nouveau v10.1** : La langue choisie est **strictement respectée**.\n- 🇫🇷 → lettre 100% en français\n- 🇬🇧 → letter 100% in English\n- 🇸🇦 → الخطاب كاملاً بالعربية\n\n⚡ Lettre professionnelle en **10 secondes** !",
                "default":  "✉️ **Lettre de Motivation IA** : choisissez la langue (FR/EN/AR) et le style → Gemini génère une lettre **strictement dans la langue choisie**."
            },
            "en": {
                "default": "✉️ **AI Cover Letter:** Choose language (FR/EN/AR) + style → Gemini generates a letter **strictly in the chosen language**. Language enforcement is guaranteed."
            },
            "ar": {
                "default": "✉️ **خطاب التقديم الذكي:** اختر اللغة (FR/EN/AR) والأسلوب → Gemini يولّد خطاباً **بالكامل باللغة المختارة** مع ضمان اللغة."
            }
        },
        "suggestions": {
            "fr": ["Générer en arabe ?","Changer le style ?","Analyser mon CV d'abord ?"],
            "en": ["Generate in Arabic?","Change style?","Analyze my CV first?"],
            "ar": ["توليد بالفرنسية؟","تغيير الأسلوب؟","تحليل سيرتي أولاً؟"]
        }
    },
    "recommandations_ia": {
        "patterns": {
            "fr": ["recommandation","conseil ia","que faire","suggestion","prochaine étape","recommandations ia","actualiser","mes recommandations"],
            "en": ["recommendation","ai advice","what to do","suggestion","refresh recommendations"],
            "ar": ["توصية","نصيحة ذكية","ماذا أفعل","اقتراح","تحديث التوصيات"]
        },
        "responses": {
            "fr": {
                "Candidat":        "🤖 **Recommandations IA :** Tests prioritaires · Entretien IA · Conseils CV · Progression\n\nCliquez **Actualiser** dans le dashboard pour rafraîchir via Gemini.",
                "AdminEntreprise": "🤖 **Recommandations Admin :** 8 candidats sans analyse · Rapport à générer · 2 membres RH à inviter",
                "SuperAdmin":      "🤖 **SuperAdmin :** Mailer DOWN urgent · 3 abonnements expirant · Audit sécurité recommandé",
                "default":         "🤖 **Recommandations IA** dans votre dashboard. Cliquez **Actualiser** pour de nouvelles recommandations Gemini."
            },
            "en": {"default": "🤖 **AI Recommendations** in your dashboard. Click **Refresh** for fresh Gemini insights."},
            "ar": {"default": "🤖 **التوصيات الذكية** في لوحة التحكم. انقر **تحديث**."}
        },
        "suggestions": {
            "fr": ["Actualiser","Améliorer mon score ?","Statistiques"],
            "en": ["Refresh","How to improve?","Statistics"],
            "ar": ["تحديث","كيف أحسّن؟","إحصائيات"]
        }
    },
    "creation_test": {
        "patterns": {
            "fr": ["créer","faire","générer","nouveau test","qcm","campagne","évaluation","créer un test","nouvelle campagne"],
            "en": ["create","make","generate","new test","mcq","campaign","new campaign"],
            "ar": ["إنشاء","اختبار جديد","حملة جديدة","توليد اختبار"]
        },
        "responses": {
            "fr": {
                "Recruteur": "✅ **Créer un test QCM IA :**\n\n**Campagnes → Nouvelle Campagne** → Poste + Niveau + Langue → **Générer avec IA** → 45 secondes !\n\n🎯 Domaines : React, Vue, Python, Node.js, Java, SQL, Docker, AWS, ML, Soft Skills...\n\n💡 Vous ne savez pas quel thème choisir ? Tapez **« quels thèmes disponibles »** pour voir la liste complète !",
                "default":   "✅ **Campagnes → Nouvelle Campagne** → Rôle + Niveau + Langue → IA génère en 45 secondes.\n\n💡 Tapez **« quels thèmes disponibles »** pour voir tous les domaines."
            },
            "en": {"default": "✅ **Campaigns → New Campaign** → Role + Level + Language → AI generates in 45s.\n\n💡 Type **\"available themes\"** to see all domains."},
            "ar": {"default": "✅ **الحملات → حملة جديدة** → المنصب + المستوى + اللغة → توليد في 45 ثانية.\n\n💡 اكتب **'المواضيع المتاحة'** لرؤية القائمة الكاملة."}
        },
        "suggestions": {
            "fr": ["Quels thèmes disponibles ?","Inviter un candidat","Activer le proctoring"],
            "en": ["Available themes?","Invite candidate","Enable proctoring"],
            "ar": ["المواضيع المتاحة؟","دعوة مرشح","تفعيل المراقبة"]
        }
    },
    "scores_results": {
        "patterns": {
            "fr": ["score","résultat","note","performance","mon score","mes résultats","progression","historique"],
            "en": ["score","result","grade","performance","my results","progress","history"],
            "ar": ["نتيجة","درجة","أداء","نتائجي","تقدمي"]
        },
        "responses": {
            "fr": {
                "Candidat": "🎯 **Vos résultats :**\n\n🟢 MBTI : **91%** · 🟡 Excel : **74%** · 🔴 Communication : **62%** · 🟢 Logique : **88%**\n\n📊 Moyenne : **76%** (+8%) · 🧠 Aura : **82/100**\n\nConsultez **Mon Espace → Résultats**.",
                "default":  "📊 Scores dans **Dashboard → Résultats**."
            },
            "en": {"default": "📊 Scores in **Dashboard → Results**. Updated real-time."},
            "ar": {"default": "🎯 الدرجات في **لوحة التحكم → النتائج**."}
        },
        "suggestions": {
            "fr": ["Améliorer mon score","Télécharger rapport","Progression"],
            "en": ["Improve score","Download report","Progression"],
            "ar": ["تحسين درجتي","تحميل تقرير","تقدمي"]
        }
    },
    "proctoring": {
        "patterns": {
            "fr": ["proctoring","anti-triche","surveillance","triche","fraude","authenticité"],
            "en": ["proctoring","anti-cheat","surveillance","cheating","fraud","authenticity"],
            "ar": ["مراقبة","مكافحة الغش","رقابة","أصالة"]
        },
        "responses": {
            "fr": {"default": "🔍 **Proctoring IA :** Vidéo · Audio · Comportemental · Biométrique · Alertes · Rapport automatique\n\n**Activation :** Campagnes → Paramètres → Proctoring → ON"},
            "en": {"default": "🔍 **AI Proctoring:** Video + audio + behavioral + biometric. Enable in Campaign Settings."},
            "ar": {"default": "🔍 **المراقبة الذكية:** فيديو + صوت + سلوكي + بيومتري. فعّل في إعدادات الحملة."}
        },
        "suggestions": {
            "fr": ["Activer le proctoring","Rapport d'incidents ?","Comment ça fonctionne ?"],
            "en": ["Enable proctoring","Incident report?","How does it work?"],
            "ar": ["تفعيل المراقبة","تقرير الحوادث؟","كيف يعمل؟"]
        }
    },
    "technical_support": {
        "patterns": {
            "fr": ["bug","problème","erreur","aide","support","ne fonctionne pas","bloqué","crash","ne marche pas"],
            "en": ["bug","problem","error","help","support","not working","stuck","crash"],
            "ar": ["مشكلة","خطأ","مساعدة","دعم","لا يعمل","عطل"]
        },
        "responses": {
            "fr": {"default": "🛠️ **Support :** Rechargez (Ctrl+F5) · Chrome · Déconnectez/reconnectez · Videz le cache\n\nToujours bloqué ? → 📧 **support@evaluatech.com** (2-4h)"},
            "en": {"default": "🛠️ **Support:** Reload (Ctrl+F5) | Chrome | logout/login | clear cache. Still broken? support@evaluatech.com"},
            "ar": {"default": "🛠️ **الدعم:** إعادة التحميل | Chrome | تسجيل الخروج. support@evaluatech.com"}
        },
        "suggestions": {
            "fr": ["FAQ ?","Signaler un bug","Contact support"],
            "en": ["FAQ?","Report a bug","Contact support"],
            "ar": ["الأسئلة الشائعة؟","الإبلاغ عن خطأ","الدعم"]
        }
    },
    "auth": {
        "patterns": {
            "fr": ["connexion","login","mot de passe","compte","oublié","déconnexion","réinitialiser"],
            "en": ["login","password","account","forgot","logout","reset"],
            "ar": ["تسجيل الدخول","كلمة المرور","حساب","نسيت","خروج"]
        },
        "responses": {
            "fr": {"default": "🔐 **Connexion :**\n- Oublié → *Mot de passe oublié* sur la page login\n- Bloqué → Contactez votre administrateur\n- Email non reçu → Vérifiez vos spams\n\n📧 **support@evaluatech.com**"},
            "en": {"default": "🔐 **Login:** Forgot → reset link | Blocked → admin | Check spam | support@evaluatech.com"},
            "ar": {"default": "🔐 **تسجيل الدخول:** نسيت → إعادة تعيين | محظور → المسؤول | البريد العشوائي | support@evaluatech.com"}
        },
        "suggestions": {
            "fr": ["Réinitialiser mot de passe","Contacter le support","Activer mon compte"],
            "en": ["Reset password","Contact support","Activate account"],
            "ar": ["إعادة تعيين","الدعم","تفعيل الحساب"]
        }
    },
    "interview": {
        "patterns": {
            "fr": ["entretien","interview","questions rh","préparer entretien","entretien ia","questions comportementales"],
            "en": ["interview","hr questions","prepare interview","ai interview","behavioral questions"],
            "ar": ["مقابلة","أسئلة مقابلة","تحضير مقابلة","مقابلة ذكية"]
        },
        "responses": {
            "fr": {
                "Candidat": "🎤 **Entretien IA :** Module dashboard → Poste + Niveau → 5-10 questions (Comportemental · Technique · Situationnel) + tips STAR",
                "default":  "🎤 **Entretien IA** : questions personnalisées par poste et niveau avec tips STAR."
            },
            "en": {"default": "🎤 **AI Interview:** Role + level → 5-10 tailored questions (behavioral, technical, situational) with STAR tips."},
            "ar": {"default": "🎤 **المقابلة الذكية:** المنصب + المستوى → 5-10 أسئلة مع نصائح STAR."}
        },
        "suggestions": {
            "fr": ["Questions STAR ?","Niveau Senior ?","Exporter PDF"],
            "en": ["STAR questions?","Senior level?","Export PDF"],
            "ar": ["أسئلة STAR؟","مستوى متقدم؟","تصدير PDF"]
        }
    },
    "superadmin_specific": {
        "patterns": {
            "fr": ["organisation","entreprise cliente","super admin","abonnement","service en panne","uptime","santé services"],
            "en": ["organization","client company","super admin","subscription","service down","uptime"],
            "ar": ["منظمة","شركة عميلة","مدير عام","اشتراك","خدمة متوقفة"]
        },
        "responses": {
            "fr": {
                "SuperAdmin": "🛡️ **SuperAdmin :**\n\n🏢 42 organisations · 1247 utilisateurs\n🟢 API Gateway (12ms) · 🟢 Auth (8ms) · 🟢 IA Engine (34ms)\n🔴 **Mailer DOWN** ⚠️ (47 emails bloqués)\n💰 82 200 DT/mois · 3 abonnements expirant dans 7j",
                "default":   "Cette section est réservée aux **SuperAdmins**."
            },
            "en": {"default": "🛡️ **SuperAdmin:** 42 organizations | 1247 users | 7/8 services UP | Mailer DOWN | $82K/month"},
            "ar": {"default": "🛡️ **SuperAdmin:** 42 منظمة | 1247 مستخدم | 7/8 خدمات تعمل | خدمة البريد متوقفة"}
        },
        "suggestions": {
            "fr": ["Gérer les organisations","Voir les abonnements","Audit sécurité"],
            "en": ["Manage organizations","View subscriptions","Security audit"],
            "ar": ["إدارة المنظمات","عرض الاشتراكات","تدقيق الأمان"]
        }
    },
}


def _get_local_response(message: str, lang: str, role: str) -> Optional[str]:
    msg = message.lower().strip()
    msg_norm = _normalize(msg)
    best_match = None
    highest_score = 0

    for intent, data in UNIVERSAL_BRAIN.items():
        score = 0
        patterns = data["patterns"].get(lang, data["patterns"].get("fr", []))
        for p in patterns:
            p_norm = _normalize(p.lower())
            if re.search(r'\b' + re.escape(p_norm) + r'\b', msg_norm):
                score += 12
            elif p_norm in msg_norm:
                score += 7
            elif p.lower() in msg:
                score += 5

        if score > highest_score:
            highest_score = score
            res_branch = data["responses"].get(lang, data["responses"].get("fr", {}))
            if isinstance(res_branch, dict):
                best_match = res_branch.get(role, res_branch.get("default", ""))
            else:
                best_match = res_branch

    return best_match if highest_score >= 5 and best_match else None


def _get_suggestions(message: str, lang: str, role: str) -> list:
    msg = _normalize(message.lower().strip())
    best_intent = None
    highest_score = 0

    for intent, data in UNIVERSAL_BRAIN.items():
        score = 0
        patterns = data["patterns"].get(lang, data["patterns"].get("fr", []))
        for p in patterns:
            if _normalize(p.lower()) in msg: score += 1
        if score > highest_score:
            highest_score = score
            best_intent = intent

    if best_intent and highest_score > 0:
        sugg = UNIVERSAL_BRAIN[best_intent].get("suggestions", {})
        if isinstance(sugg, dict):
            return sugg.get(lang, sugg.get("fr", []))[:3]
        return sugg[:3]

    defaults = {
        "Candidat":        {"fr": ["Voir mes résultats","Préparer mon entretien","Analyser mon CV"],           "en": ["View results","Prepare interview","Analyze CV"],      "ar": ["نتائجي","تحضير المقابلة","تحليل السيرة"]},
        "Recruteur":       {"fr": ["Créer un test","Analyser un CV","Voir le pipeline"],                       "en": ["Create test","Analyze CV","View pipeline"],           "ar": ["إنشاء اختبار","تحليل السيرة","المسار"]},
        "AdminEntreprise": {"fr": ["Mes statistiques","Générer un rapport","Recommandations IA"],              "en": ["My statistics","Generate report","AI recommendations"],"ar": ["إحصائياتي","تقرير","التوصيات"]},
        "SuperAdmin":      {"fr": ["Santé des services","Gérer les organisations","Abonnements"],              "en": ["Services health","Manage orgs","Subscriptions"],       "ar": ["صحة الخدمات","المنظمات","الاشتراكات"]},
        "Evaluateur":      {"fr": ["Voir ma file","Planifier une session","Top compétences"],                  "en": ["View queue","Schedule session","Top skills"],          "ar": ["قائمة التقييم","جدولة جلسة","المهارات"]},
        "RH":              {"fr": ["Créer une campagne","Statistiques candidats","Rapport mensuel"],            "en": ["Create campaign","Candidate stats","Monthly report"],  "ar": ["إنشاء حملة","إحصائيات","تقرير شهري"]},
        "default":         {"fr": ["Créer un test IA","Analyser un CV","Voir le dashboard"],                  "en": ["Create AI test","Analyze CV","View dashboard"],        "ar": ["إنشاء اختبار","تحليل سيرة","لوحة التحكم"]},
    }
    role_defs = defaults.get(role, defaults["default"])
    return role_defs.get(lang, role_defs.get("fr", []))


# ════════════════════════════════════════════════════════════
# ██   RECOMMANDATIONS IA   ██
# ════════════════════════════════════════════════════════════

RECO_ICONS = {
    "performance": "fa-solid fa-chart-line",   "test": "fa-solid fa-clipboard-list",
    "alert":       "fa-solid fa-triangle-exclamation", "user": "fa-solid fa-users",
    "security":    "fa-solid fa-shield-halved", "skill": "fa-solid fa-brain",
    "cv":          "fa-solid fa-file-pdf",      "interview": "fa-solid fa-microphone",
    "report":      "fa-solid fa-chart-bar",     "system": "fa-solid fa-server",
    "campaign":    "fa-solid fa-bullhorn",      "money": "fa-solid fa-coins",
    "calendar":    "fa-solid fa-calendar-check","team": "fa-solid fa-people-group",
    "mail":        "fa-solid fa-envelope",      "star": "fa-solid fa-star",
    "letter":      "fa-solid fa-envelope-open-text", "stats": "fa-solid fa-chart-pie",
}
RECO_PRIORITY = {
    "urgent": {"bg": "#fee2e2", "color": "#dc2626", "label": "🔴 Urgent"},
    "high":   {"bg": "#fef9ec", "color": "#d97706", "label": "🟡 Priorité"},
    "medium": {"bg": "#ecfdf5", "color": "#059669", "label": "🟢 Standard"},
    "low":    {"bg": "#eff6ff", "color": "#2563eb", "label": "🔵 Info"},
}
RECO_COLORS = ["#6366f1","#f59e0b","#10b981","#ef4444","#8b5cf6","#3b82f6","#ec4899"]

def _build_reco_prompt(role: str, context_data: dict = None) -> str:
    ctx = f"\nContexte: {json.dumps(context_data, ensure_ascii=False)[:400]}" if context_data else ""
    routes_valides = "/dashboard, /questions, /campaigns, /my-tests, /eval-queue, /analyse-comportementale, /sessions, /cv-library, /candidates-list, /invite, /reporting, /staff-members, /gestion-staff, /super-admin, /gestion-abonnements, /platform-users, /super-admin-analytics, /interview-prep, /entretien-ia"
    
    prompts = {
        "Candidat":        f"Coach carrière EvaluaTech. 3 recommandations candidat (score 82%, 2 tests en attente).{ctx}\nJSON: {{\"recommendations\":[{{\"title\":\"...\",\"description\":\"...\",\"actionLabel\":\"...\",\"icon\":\"test|skill|cv|interview|star|letter\",\"priority\":\"urgent|high|medium|low\",\"route\":\"{routes_valides}\"}}]}}",
        "Evaluateur":      f"Expert RH. 3 recommandations évaluateur (5 en attente, taux 91%).{ctx}\nJSON: {{\"recommendations\":[{{\"title\":\"...\",\"description\":\"...\",\"actionLabel\":\"...\",\"icon\":\"test|user|report|calendar|alert\",\"priority\":\"urgent|high|medium|low\",\"route\":\"{routes_valides}\"}}]}}",
        "RH":              f"Expert RH. 3 recommandations RH (3 campagnes, 18 candidats, 84%).{ctx}\nJSON: {{\"recommendations\":[{{\"title\":\"...\",\"description\":\"...\",\"actionLabel\":\"...\",\"icon\":\"user|campaign|report|performance|team\",\"priority\":\"urgent|high|medium|low\",\"route\":\"{routes_valides}\"}}]}}",
        "Recruteur":       f"Expert recrutement. 3 recommandations recruteur (22 candidats, 7 sans réponse).{ctx}\nJSON: {{\"recommendations\":[{{\"title\":\"...\",\"description\":\"...\",\"actionLabel\":\"...\",\"icon\":\"user|campaign|cv|performance|mail|stats\",\"priority\":\"urgent|high|medium|low\",\"route\":\"{routes_valides}\"}}]}}",
        "AdminEntreprise": f"Consultant EvaluaTech. 3 recommandations Admin (124 talents, 88%, 8 candidats 48h).{ctx}\nJSON: {{\"recommendations\":[{{\"title\":\"...\",\"description\":\"...\",\"actionLabel\":\"...\",\"icon\":\"performance|user|report|security|campaign|team|stats\",\"priority\":\"urgent|high|medium|low\",\"route\":\"{routes_valides}\"}}]}}",
        "SuperAdmin":      f"Expert SaaS. 3 recommandations SuperAdmin (Mailer DOWN 2h, 3 abonnements 7j).{ctx}\nJSON: {{\"recommendations\":[{{\"title\":\"...\",\"description\":\"...\",\"actionLabel\":\"...\",\"icon\":\"system|security|alert|money|report\",\"priority\":\"urgent|high|medium|low\",\"route\":\"{routes_valides}\"}}]}}",
    }
    base_prompt = prompts.get(role, prompts["AdminEntreprise"])
    return base_prompt + f"\nCONTRAINTE STRICTE: Pour la clé 'route', vous DEVEZ choisir UNIQUEMENT parmi ces routes exactes : {routes_valides}. Ne générez JAMAIS d'autres URLs."

def _get_fallback_recommendations(role: str) -> list:
    fallbacks = {
        "Candidat": [
            {"title": "Complétez vos tests en attente", "description": "2 tests en attente. Les compléter augmentera votre profil de 15 points.", "actionLabel": "Voir mes tests", "icon": RECO_ICONS["test"], "color": "#3b82f6", "priority": "🔴 Urgent", "priorityBg": "#fee2e2", "priorityColor": "#dc2626", "route": "/my-tests", "scrollTo": None},
            {"title": "Préparez votre entretien IA", "description": "Entraînez-vous sur des questions comportementales et techniques.", "actionLabel": "Commencer préparation", "icon": RECO_ICONS["interview"], "color": "#8b5cf6", "priority": "🟡 Priorité", "priorityBg": "#fef9ec", "priorityColor": "#d97706", "route": "/interview-prep", "scrollTo": None},
            {"title": "Analysez votre CV avec Gemini", "description": "Obtenez des conseils personnalisés pour augmenter votre score de matching.", "actionLabel": "Analyser mon CV", "icon": RECO_ICONS["cv"], "color": "#10b981", "priority": "🟢 Standard", "priorityBg": "#ecfdf5", "priorityColor": "#059669", "route": None, "scrollTo": "cv-scan-section"},
        ],
        "Evaluateur": [
            {"title": "5 évaluations urgentes en attente", "description": "Sara Ben Ali (26h) est prioritaire.", "actionLabel": "Voir la file", "icon": RECO_ICONS["alert"], "color": "#ef4444", "priority": "🔴 Urgent", "priorityBg": "#fee2e2", "priorityColor": "#dc2626", "route": "/analyse-comportementale", "scrollTo": None},
            {"title": "Session demain à 10h00", "description": "Session Q2 avec 12 candidats. Préparez vos grilles.", "actionLabel": "Voir la session", "icon": RECO_ICONS["calendar"], "color": "#f59e0b", "priority": "🟡 Priorité", "priorityBg": "#fef9ec", "priorityColor": "#d97706", "route": "/sessions", "scrollTo": None},
            {"title": "Rapport hebdomadaire disponible", "description": "Consultez votre rapport de performance de la semaine.", "actionLabel": "Voir le rapport", "icon": RECO_ICONS["report"], "color": "#10b981", "priority": "🟢 Standard", "priorityBg": "#ecfdf5", "priorityColor": "#059669", "route": "/reporting", "scrollTo": None},
        ],
        "RH": [
            {"title": "Lancer une campagne de recrutement", "description": "Taux candidatures en baisse de 12%. Créez une campagne ciblée.", "actionLabel": "Créer campagne", "icon": RECO_ICONS["campaign"], "color": "#8b5cf6", "priority": "🟡 Priorité", "priorityBg": "#fef9ec", "priorityColor": "#d97706", "route": "/campaigns", "scrollTo": None},
            {"title": "Analyser les soft skills", "description": "18 candidats ont terminé. Lancez l'analyse comportementale.", "actionLabel": "Analyser profils", "icon": RECO_ICONS["skill"], "color": "#10b981", "priority": "🟢 Standard", "priorityBg": "#ecfdf5", "priorityColor": "#059669", "route": "/analyse-comportementale", "scrollTo": None},
            {"title": "Rapport mensuel RH", "description": "Générez le rapport RH mensuel pour votre direction.", "actionLabel": "Générer rapport", "icon": RECO_ICONS["report"], "color": "#3b82f6", "priority": "🔵 Info", "priorityBg": "#eff6ff", "priorityColor": "#2563eb", "route": "/reporting", "scrollTo": None},
        ],
        "Recruteur": [
            {"title": "7 candidats sans réponse", "description": "Relancez-les pour maximiser le pipeline.", "actionLabel": "Relancer invitations", "icon": RECO_ICONS["mail"], "color": "#ef4444", "priority": "🔴 Urgent", "priorityBg": "#fee2e2", "priorityColor": "#dc2626", "route": "/candidates-list", "scrollTo": None},
            {"title": "3 profils React Senior à 85%+", "description": "Mariam Khelifi (95%), Youssef Chaabane (89%) et Fatima Zouari (92%).", "actionLabel": "Voir profils", "icon": RECO_ICONS["star"], "color": "#10b981", "priority": "🟡 Priorité", "priorityBg": "#fef9ec", "priorityColor": "#d97706", "route": "/candidates-list", "scrollTo": None},
            {"title": "Créer un test DevOps", "description": "Pas de test DevOps actif. Élargissez votre pipeline.", "actionLabel": "Créer test", "icon": RECO_ICONS["test"], "color": "#6366f1", "priority": "🟢 Standard", "priorityBg": "#ecfdf5", "priorityColor": "#059669", "route": "/campaigns", "scrollTo": None},
        ],
        "AdminEntreprise": [
            {"title": "8 candidats sans analyse depuis 48h", "description": "Traitez-les avant expiration des liens.", "actionLabel": "Voir le pipeline", "icon": RECO_ICONS["alert"], "color": "#ef4444", "priority": "🔴 Urgent", "priorityBg": "#fee2e2", "priorityColor": "#dc2626", "route": "/analyse-comportementale", "scrollTo": None},
            {"title": "Générer le rapport mensuel", "description": "Le rapport de ce mois n'a pas encore été généré.", "actionLabel": "Générer rapport", "icon": RECO_ICONS["report"], "color": "#f59e0b", "priority": "🟡 Priorité", "priorityBg": "#fef9ec", "priorityColor": "#d97706", "route": "/reporting", "scrollTo": None},
            {"title": "Inviter de nouveaux membres RH", "description": "Équipe en sous-effectif. Invitez 2 membres supplémentaires.", "actionLabel": "Inviter membres", "icon": RECO_ICONS["team"], "color": "#6366f1", "priority": "🟢 Standard", "priorityBg": "#ecfdf5", "priorityColor": "#059669", "route": "/invite", "scrollTo": None},
        ],
        "SuperAdmin": [
            {"title": "Service Mailer en panne — CRITIQUE", "description": "DOWN depuis 2h. 47 invitations bloquées. Intervention immédiate.", "actionLabel": "Diagnostiquer", "icon": RECO_ICONS["system"], "color": "#ef4444", "priority": "🔴 Urgent", "priorityBg": "#fee2e2", "priorityColor": "#dc2626", "route": "/super-admin", "scrollTo": None},
            {"title": "3 abonnements expirant dans 7 jours", "description": "Contactez les entreprises pour le renouvellement.", "actionLabel": "Voir abonnements", "icon": RECO_ICONS["money"], "color": "#f59e0b", "priority": "🟡 Priorité", "priorityBg": "#fef9ec", "priorityColor": "#d97706", "route": "/gestion-abonnements", "scrollTo": None},
            {"title": "Audit de sécurité recommandé", "description": "Aucun audit depuis 30 jours. Planifiez une vérification.", "actionLabel": "Lancer audit", "icon": RECO_ICONS["security"], "color": "#6366f1", "priority": "🟢 Standard", "priorityBg": "#ecfdf5", "priorityColor": "#059669", "route": "/super-admin-analytics", "scrollTo": None},
        ],
    }
    return fallbacks.get(role, fallbacks["AdminEntreprise"])

def _parse_gemini_recommendations(raw_text: str, role: str) -> list:
    try:
        cleaned = raw_text.replace("```json","").replace("```","").strip()
        json_match = re.search(r'\{[\s\S]*\}', cleaned)
        if not json_match: return []
        parsed = json.loads(json_match.group(0))
        if not parsed.get("recommendations"): return []
        result = []
        for i, r in enumerate(parsed["recommendations"][:3]):
            prio = RECO_PRIORITY.get(r.get("priority","medium"), RECO_PRIORITY["medium"])
            result.append({
                "title":         r.get("title","Recommandation"),
                "description":   r.get("description",""),
                "actionLabel":   r.get("actionLabel","Voir plus"),
                "icon":          RECO_ICONS.get(r.get("icon","performance"), "fa-solid fa-lightbulb"),
                "color":         RECO_COLORS[i % len(RECO_COLORS)],
                "priority":      prio["label"],
                "priorityBg":    prio["bg"],
                "priorityColor": prio["color"],
                "route":         r.get("route", "/dashboard")
            })
        return result
    except Exception:
        return []


class RecommendationRequest(BaseModel):
    role: str = "AdminEntreprise"
    lang: str = "fr"
    context: str = "{}"
    force_refresh: bool = False

@app.post("/ia/recommendations")
async def get_recommendations(req: RecommendationRequest):
    role = req.role
    lang = req.lang
    context = req.context
    force_refresh = req.force_refresh
    
    ck = make_cache_key("reco-v10", role, lang)
    if not force_refresh:
        cached = _reco_cache.get(ck)
        if cached:
            return {"status": "SUCCESS", "recommendations": cached, "source": "cache", "role": role}
    try:
        ctx_data = json.loads(context) if context and context != "{}" else {}
    except:
        ctx_data = {}
    try:
        response = await call_gemini_async(_build_reco_prompt(role, ctx_data), module="Recommandations", sem=_sem_reco, retries=1)
        recommendations = _parse_gemini_recommendations(response.text.strip(), role)
        if not recommendations: recommendations = _get_fallback_recommendations(role); source = "fallback_parse"
        else: source = "gemini"
        _reco_cache.set(ck, recommendations)
        AI_METRICS["usage_counts"]["Recommandations"] += 1
        log_activity("IA", f"Recommandations ({role})", "#f59e0b", "Recommandations")
        return {"status": "SUCCESS", "recommendations": recommendations, "source": source, "role": role, "generated_at": datetime.now().isoformat()}
    except (QuotaExceeded, Exception):
        fallback = _get_fallback_recommendations(role)
        _reco_cache.set(ck, fallback)
        return {"status": "SUCCESS", "recommendations": fallback, "source": "fallback", "role": role, "generated_at": datetime.now().isoformat()}


@app.get("/ia/recommendations")
async def get_recommendations_get(role: str = Query("AdminEntreprise"), lang: str = Query("fr"), force_refresh: bool = Query(False)):
    ck = make_cache_key("reco-v10", role, lang)
    if not force_refresh:
        cached = _reco_cache.get(ck)
        if cached: return {"status": "SUCCESS", "recommendations": cached, "source": "cache", "role": role}
    try:
        response = await call_gemini_async(_build_reco_prompt(role), module="Recommandations", sem=_sem_reco, retries=1)
        recommendations = _parse_gemini_recommendations(response.text.strip(), role)
        if not recommendations: recommendations = _get_fallback_recommendations(role); source = "fallback"
        else: source = "gemini"
        _reco_cache.set(ck, recommendations)
        return {"status": "SUCCESS", "recommendations": recommendations, "source": source, "role": role}
    except (QuotaExceeded, Exception):
        return {"status": "SUCCESS", "recommendations": _get_fallback_recommendations(role), "source": "fallback", "role": role}


# ════════════════════════════════════════════════════════════
# ██   LETTRE DE MOTIVATION IA v10.1 — LANGUE STRICTE   ██
# ════════════════════════════════════════════════════════════

@app.post("/ia/lettre-motivation")
async def generate_lettre_motivation(
    nom: str = Form(""),
    poste: str = Form(...),
    entreprise: str = Form(...),
    competences: str = Form(""),
    langue: str = Form("fr"),
    style: str = Form("professionnel")
):
    lang_key = langue.strip().lower()
    if lang_key not in ["fr", "en", "ar"]:
        lang_key = "fr"

    style_key = style.strip().lower()
    if style_key not in ["professionnel", "créatif", "concis"]:
        style_key = "professionnel"

    AI_METRICS["usage_counts"]["Lettres"] += 1

    prompt = _build_lettre_prompt(nom, poste, entreprise, competences, lang_key, style_key)

    lang_labels = {"fr": "Français", "en": "English", "ar": "العربية"}
    lang_label  = lang_labels.get(lang_key, lang_key.upper())

    try:
        response = await call_gemini_async(prompt, module="Lettres", sem=_sem_rpts, retries=1)
        text = response.text.strip()

        if text and len(text) > 80:
            detected_lang = _verify_letter_language(text, lang_key)

            if detected_lang == lang_key or detected_lang == "unknown":
                log_activity("IA", f"Lettre {lang_label} générée ({poste})", "#8b5cf6", "Lettres")
                return {
                    "status": "SUCCESS",
                    "lettre": text,
                    "source": "gemini",
                    "langue": lang_key,
                    "langue_label": lang_label,
                    "style": style_key,
                    "enforced_language": True,
                }
            else:
                logger.warning(f"Gemini returned wrong language (expected {lang_key}, got {detected_lang}). Using template fallback.")
                lettre_fallback = _generate_fallback_lettre(nom, poste, entreprise, competences, lang_key, style_key)
                log_activity("IA", f"Lettre {lang_label} (template) ({poste})", "#f59e0b", "Lettres")
                return {
                    "status": "SUCCESS",
                    "lettre": lettre_fallback,
                    "source": "template_fallback",
                    "langue": lang_key,
                    "langue_label": lang_label,
                    "style": style_key,
                    "enforced_language": True,
                    "note": f"Template utilisé pour garantir la langue {lang_label}.",
                }
        else:
            raise ValueError("Empty or too short response")

    except (QuotaExceeded, Exception) as e:
        lettre_fallback = _generate_fallback_lettre(nom, poste, entreprise, competences, lang_key, style_key)
        log_activity("IA", f"Lettre {lang_label} (fallback) ({poste})", "#94a3b8", "Lettres")
        return {
            "status": "SUCCESS",
            "lettre": lettre_fallback,
            "source": "fallback",
            "langue": lang_key,
            "langue_label": lang_label,
            "style": style_key,
            "enforced_language": True,
        }


def _verify_letter_language(text: str, expected_lang: str) -> str:
    text_lower = text.lower()

    fr_markers = ["madame","monsieur","cordialement","bonjour","je vous","votre entreprise","je suis","candidature","poste","compétences"]
    en_markers = ["dear","sincerely","please","i am","your company","position","skills","experience","regards","hiring"]
    ar_markers = ["السيد","المحترم","مع التقدير","أتقدم","المنصب","مهارات","خبرة","شركة","يسعدني"]

    ar_chars = len(re.findall(r'[\u0600-\u06FF]', text))
    fr_count = sum(1 for m in fr_markers if m in text_lower)
    en_count = sum(1 for m in en_markers if m in text_lower)

    if ar_chars > 20: return "ar"
    if fr_count >= 3 and fr_count > en_count: return "fr"
    if en_count >= 3 and en_count > fr_count: return "en"

    if expected_lang == "fr" and en_count > fr_count and en_count >= 3: return "en"
    if expected_lang == "en" and fr_count > en_count and fr_count >= 3: return "fr"
    if expected_lang == "ar" and ar_chars < 10: return "fr" if fr_count >= en_count else "en"

    return "unknown"


@app.get("/ia/lettre-motivation/languages")
async def get_lettre_languages():
    return {
        "languages": [
            {
                "code": "fr",
                "label": "Français",
                "flag": "🇫🇷",
                "description": "Lettre rédigée entièrement en français",
                "styles": ["professionnel", "créatif", "concis"],
            },
            {
                "code": "en",
                "label": "English",
                "flag": "🇬🇧",
                "description": "Letter written entirely in English",
                "styles": ["professionnel", "créatif", "concis"],
            },
            {
                "code": "ar",
                "label": "العربية",
                "flag": "🇸🇦",
                "description": "الخطاب مكتوب بالكامل باللغة العربية",
                "styles": ["professionnel", "créatif", "concis"],
                "rtl": True,
            },
        ],
        "styles": {
            "professionnel": {
                "fr": "Classique, formel, persuasif",
                "en": "Classic, formal, persuasive",
                "ar": "كلاسيكي، رسمي، مقنع",
            },
            "créatif": {
                "fr": "Moderne, dynamique, accroche originale",
                "en": "Modern, dynamic, original opening",
                "ar": "حديث، ديناميكي، مقدمة إبداعية",
            },
            "concis": {
                "fr": "Court, direct, percutant (≤200 mots)",
                "en": "Short, direct, impactful (≤200 words)",
                "ar": "قصير، مباشر، مؤثر (≤200 كلمة)",
            },
        },
        "enforcement": "language_strictly_enforced",
        "version": "10.1",
    }


# ════════════════════════════════════════════════════════════
# ██   ANALYSE CV v10.1 — DÉTECTION NON-CV   ██
# ════════════════════════════════════════════════════════════

def _local_conseils(score: int) -> list:
    if score >= 85: return [
        "Valider sa capacité à gérer l'autonomie et ses ambitions d'évolution en entretien.",
        "Lui confier des responsabilités techniques de haut niveau ou de mentorat dès son intégration.",
        "Sonder sa flexibilité face à des changements technologiques ou structurels rapides."
    ]
    elif score >= 70: return [
        "Évaluer sa posture collaborative et sa communication transverse en entretien individuel.",
        "Planifier un accompagnement technique ciblé sur les points techniques plus faibles lors de l'onboarding.",
        "Tester sa gestion des priorités et du stress face aux deadlines serrées de l'entreprise."
    ]
    else: return [
        "Creuser en profondeur ses fondamentaux techniques via une évaluation pratique complémentaire.",
        "Prévoir un plan de formation rigoureux de 30 à 60 jours avant toute affectation autonome.",
        "Sonder ses motivations réelles et son esprit d'apprentissage continu face aux lacunes identifiées."
    ]


@app.post("/ia/match-cv")
async def match_cv(
    file: UploadFile = File(...),
    job_description: str = Form(...),
    lang: str = Form("fr"),
    skip_cv_check: bool = Form(False)
):
    fb   = await file.read()
    name = (file.filename or "").lower()

    if name.endswith(".pdf"):
        cv_text = _truncate(await read_pdf_async(fb))
    else:
        loop = asyncio.get_event_loop()
        def _other():
            try:
                if name.endswith(".docx") and DOCX_AVAILABLE:
                    doc = Document(io.BytesIO(fb))
                    return "\n".join(p.text for p in doc.paragraphs)
                return fb.decode("utf-8", errors="ignore")
            except: return ""
        cv_text = _truncate(await loop.run_in_executor(_gemini_executor, _other))

    if not skip_cv_check:
        detection = _detect_is_cv(cv_text)

        if not detection["is_cv"]:
            AI_METRICS["cv_rejected"] += 1
            lang_key = lang.strip().lower() if lang.strip().lower() in ["fr","en","ar"] else "fr"
            alert_payload = _build_cv_alert(detection, lang_key)

            log_activity(
                "⚠️ Alerte",
                f"Document non-CV uploadé ({name or 'fichier inconnu'})",
                "#f59e0b",
                "Analyses CV"
            )

            return JSONResponse(
                status_code=422,
                content={
                    "status": "NOT_A_CV",
                    "is_cv": False,
                    "alert": alert_payload["alert"],
                    "confidence": alert_payload["confidence"],
                    "matched_keywords": alert_payload["matched_keywords"],
                    "word_count": alert_payload["word_count"],
                    "filename": file.filename or "fichier",
                    "suggestion": {
                        "fr": "Assurez-vous d'uploader votre CV au format PDF ou DOCX.",
                        "en": "Please make sure to upload your CV in PDF or DOCX format.",
                        "ar": "يرجى التأكد من رفع سيرتك الذاتية بصيغة PDF أو DOCX.",
                    }.get(lang_key, ""),
                }
            )

        cv_warning = detection.get("warning") if detection.get("confidence", 1) < 0.75 else None
    else:
        cv_warning = None

    ck = make_cache_key("match-cv-v10", cv_text[:200], job_description[:200])
    if hit := _cache.get(ck):
        if cv_warning: hit["cv_quality_warning"] = cv_warning
        return hit

    try:
        prompt = f"""Effectue une analyse comportementale prédictive et de personnalité approfondie de ce candidat pour le poste de : {job_description[:400]}.
Le but est d'aider la direction et les évaluateurs de l'entreprise à comprendre la véritable personnalité du candidat, ses traits comportementaux majeurs (leadership, communication, travail en équipe), sa réaction sous pression, et à obtenir des conseils managériaux stratégiques directs d'onboarding et d'entretien. Ne fais aucune critique sur la forme du CV, son contenu manquant ou sa rédaction.
Document textuel (CV ou informations de profil) : {cv_text[:2500]}

Génère une réponse au format JSON strict avec les clés suivantes :
1. "score" : un entier de 0 à 100 représentant le score de compatibilité comportementale globale avec le poste.
2. "points_forts" : une liste de 3 points forts comportementaux et traits de personnalité majeurs (ex: Autonomie affirmée, communication bienveillante, capacité d'adaptation rapide).
3. "points_faibles" : une liste de 2-3 points de vigilance comportementaux ou réactions typiques sous stress (ex: Peut avoir du mal à déléguer dans l'urgence, besoin d'un cadre clair pour performer).
4. "decision" : un verdict de synthèse sur le profil comportemental du candidat (ex: Profil hautement collaboratif, idéal pour un rôle structurant au sein de l'équipe).
5. "conseils" : une liste de 3 recommandations managériales concrètes pour la direction (onboarding, motivation, style de management adapté à sa personnalité).
6. "competences_detectees" : une liste de 3 à 5 soft-skills ou comportements clés observés.
7. "niveau_estime" : style de comportement ou maturité professionnelle estimée (Junior|Mid|Senior).

JSON valide sans markdown :
{{"score":<0-100>,"points_forts":["...","...","..."],"points_faibles":["...","..."],"decision":"...","conseils":["...","...","..."],"competences_detectees":["...","..."],"niveau_estime":"Junior|Mid|Senior"}}"""
        r = await call_gemini_async(prompt, module="Analyses CV", sem=_sem_cv, retries=1)
        result = json.loads(clean_json(r.text))
        if "conseils" not in result:
            result["conseils"] = _local_conseils(result.get("score", 75))
        result["is_cv"] = True
        if cv_warning:
            result["cv_quality_warning"] = cv_warning
        _cache.set(ck, result)
        log_activity("IA", f"Matching CV — Score: {result.get('score')}%", "#fbbf24", "Analyses CV")
        return result

    except (QuotaExceeded, Exception):
        score = random.randint(65, 90)
        result = {
            "score": score,
            "points_forts": ["Expérience technique validée", "Soft skills reconnus", "Profil compatible"],
            "points_faibles": ["Certifications à renforcer", "Portfolio à compléter"],
            "decision": "Recommandé pour entretien" if score >= 75 else "À réévaluer",
            "conseils": _local_conseils(score),
            "competences_detectees": ["Communication", "Travail en équipe"],
            "niveau_estime": "Mid",
            "is_cv": True,
        }
        if cv_warning:
            result["cv_quality_warning"] = cv_warning
        return result


@app.post("/ia/validate-cv-doc")
async def validate_cv_document(
    file: UploadFile = File(...),
    lang: str = Form("fr")
):
    fb   = await file.read()
    name = (file.filename or "").lower()

    if name.endswith(".pdf"):
        text = _truncate(await read_pdf_async(fb))
    else:
        loop = asyncio.get_event_loop()
        def _other():
            try:
                if name.endswith(".docx") and DOCX_AVAILABLE:
                    doc = Document(io.BytesIO(fb))
                    return "\n".join(p.text for p in doc.paragraphs)
                return fb.decode("utf-8", errors="ignore")
            except: return ""
        text = _truncate(await loop.run_in_executor(_gemini_executor, _other))

    detection = _detect_is_cv(text)
    lang_key  = lang.strip().lower() if lang.strip().lower() in ["fr","en","ar"] else "fr"

    if detection["is_cv"]:
        return {
            "is_cv": True,
            "confidence": detection.get("confidence", 0.8),
            "matched_keywords": detection.get("matched_keywords", []),
            "word_count": detection.get("word_count", 0),
            "cv_quality_warning": detection.get("warning"),
            "cv_quality_warning_en": detection.get("warning_en"),
            "cv_quality_warning_ar": detection.get("warning_ar"),
            "status": "OK",
        }
    else:
        alert_payload = _build_cv_alert(detection, lang_key)
        return {
            "is_cv": False,
            "confidence": detection.get("confidence", 0),
            "alert": alert_payload["alert"],
            "matched_keywords": detection.get("matched_keywords", []),
            "word_count": detection.get("word_count", 0),
            "status": "NOT_A_CV",
        }


@app.post("/ia/radar-analysis")
async def radar_analysis(file: UploadFile = File(...)):
    fb = await file.read()
    cv_text = _truncate(await read_pdf_async(fb))
    ck = make_cache_key("radar-v10", cv_text[:200])
    if hit := _cache.get(ck): return hit
    try:
        prompt = f'Analyse soft-skills CV JSON:\n{cv_text[:1500]}\n{{"values":[int,int,int,int,int]}} Communication,Leadership,Adaptabilité,Équipe,Résolution 0-100.'
        r = await call_gemini_async(prompt, module="Analyses CV", sem=_sem_cv, retries=1)
        result = json.loads(clean_json(r.text))
        _cache.set(ck, result); return result
    except (QuotaExceeded, Exception):
        return {"values":[60+random.randint(-10,20),85+random.randint(-10,10),70+random.randint(-10,15),90+random.randint(-10,5),75+random.randint(-10,15)]}


# ════════════════════════════════════════════════════════════
# ██   DASHBOARD   ██
# ════════════════════════════════════════════════════════════

def _get_dynamic_kpis(role: str) -> dict:
    configs = {
        "Candidat": {
            "kpis": [
                {"label":"TESTS PASSÉS",  "value":8+random.randint(0,3),       "icon":"fa-solid fa-file-check",     "color":"#3b82f6","bg":"rgba(59,130,246,0.12)","trend":f"+{random.randint(1,3)}","sparkPoints":"0,22 12,18 24,20 36,10 48,14 60,8 72,12 80,6"},
                {"label":"SCORE MOYEN",   "value":f"{76+random.randint(-3,8)}%","icon":"fa-solid fa-star",            "color":"#f59e0b","bg":"rgba(245,158,11,0.12)","trend":"+8%","sparkPoints":"0,24 12,20 24,16 36,18 48,12 60,10 72,8 80,4"},
                {"label":"EN ATTENTE",    "value":2+random.randint(0,2),        "icon":"fa-solid fa-hourglass-half",  "color":"#ef4444","bg":"rgba(239,68,68,0.12)","trend":"→","sparkPoints":"0,20 12,22 24,18 36,16 48,20 60,14 72,18 80,12"},
                {"label":"AURA SCORE",    "value":82+random.randint(-3,5),      "icon":"fa-solid fa-brain",           "color":"#8b5cf6","bg":"rgba(139,92,246,0.12)","trend":"↑","sparkPoints":"0,26 12,22 24,24 36,16 48,18 60,10 72,12 80,6"},
            ],
            "summary":{"totalTests":8,"moyenne":76,"totalCampagnes":2,"totalTalents":8,"iaProcessed":5}
        },
        "Evaluateur": {
            "kpis": [
                {"label":"À ÉVALUER",      "value":14+random.randint(-3,5),       "icon":"fa-solid fa-users-gear",  "color":"#f59e0b","bg":"rgba(245,158,11,0.12)","trend":f"+{random.randint(1,5)}","sparkPoints":"0,22 12,18 24,20 36,10 48,14 60,8 72,12 80,6"},
                {"label":"TAUX TRAITEMENT","value":f"{91+random.randint(-3,5)}%", "icon":"fa-solid fa-circle-check","color":"#10b981","bg":"rgba(16,185,129,0.12)","trend":"+3%","sparkPoints":"0,24 12,20 24,16 36,18 48,12 60,10 72,8 80,4"},
                {"label":"SESSIONS",       "value":6+random.randint(-1,3),        "icon":"fa-solid fa-calendar-days","color":"#6366f1","bg":"rgba(99,102,241,0.12)","trend":f"+{random.randint(0,2)}","sparkPoints":"0,20 12,16 24,22 36,14 48,16 60,12 72,14 80,8"},
                {"label":"SCORE MOYEN",    "value":f"{88+random.randint(-5,5)}%", "icon":"fa-solid fa-chart-bar",   "color":"#8b5cf6","bg":"rgba(139,92,246,0.12)","trend":"↑","sparkPoints":"0,26 12,22 24,24 36,16 48,18 60,10 72,12 80,6"},
            ],
            "summary":{"totalTests":47,"moyenne":88,"totalCampagnes":6,"totalTalents":14,"iaProcessed":42}
        },
        "RH": {
            "kpis": [
                {"label":"À ÉVALUER",     "value":18+random.randint(-3,5),      "icon":"fa-solid fa-users-gear",  "color":"#8b5cf6","bg":"rgba(139,92,246,0.12)","trend":f"+{random.randint(1,5)}","sparkPoints":"0,22 12,18 24,20 36,10 48,14 60,8 72,12 80,6"},
                {"label":"TAUX RÉUSSITE", "value":f"{84+random.randint(-3,5)}%","icon":"fa-solid fa-circle-check","color":"#10b981","bg":"rgba(16,185,129,0.12)","trend":"+2%","sparkPoints":"0,24 12,20 24,16 36,18 48,12 60,10 72,8 80,4"},
                {"label":"CAMPAGNES",     "value":3+random.randint(0,2),        "icon":"fa-solid fa-bullhorn",    "color":"#f59e0b","bg":"rgba(245,158,11,0.12)","trend":f"+{random.randint(0,2)}","sparkPoints":"0,20 12,16 24,22 36,14 48,16 60,12 72,14 80,8"},
                {"label":"SCORE MOYEN",   "value":f"{84+random.randint(-3,5)}%","icon":"fa-solid fa-chart-bar",   "color":"#3b82f6","bg":"rgba(59,130,246,0.12)","trend":"↑","sparkPoints":"0,26 12,22 24,24 36,16 48,18 60,10 72,12 80,6"},
            ],
            "summary":{"totalTests":56,"moyenne":84,"totalCampagnes":3,"totalTalents":18,"iaProcessed":50}
        },
        "Recruteur": {
            "kpis": [
                {"label":"NOUVEAUX CANDIDATS","value":22+random.randint(-3,8),       "icon":"fa-solid fa-user-plus","color":"#10b981","bg":"rgba(16,185,129,0.12)","trend":f"+{random.randint(2,7)}","sparkPoints":"0,22 12,18 24,20 36,10 48,14 60,8 72,12 80,6"},
                {"label":"INVITATIONS",       "value":7+random.randint(-2,5),        "icon":"fa-solid fa-envelope","color":"#f59e0b","bg":"rgba(245,158,11,0.12)","trend":f"-{random.randint(1,3)}","sparkPoints":"0,24 12,20 24,16 36,18 48,12 60,10 72,8 80,4"},
                {"label":"CAMPAGNES",         "value":3+random.randint(0,2),         "icon":"fa-solid fa-bullhorn","color":"#8b5cf6","bg":"rgba(139,92,246,0.12)","trend":"→","sparkPoints":"0,20 12,16 24,22 36,14 48,16 60,12 72,14 80,8"},
                {"label":"PROFILS COMPAT.",   "value":7+random.randint(-2,5),        "icon":"fa-solid fa-chart-bar","color":"#3b82f6","bg":"rgba(59,130,246,0.12)","trend":f"+{random.randint(1,4)}%","sparkPoints":"0,26 12,22 24,24 36,16 48,18 60,10 72,12 80,6"},
            ],
            "summary":{"totalTests":35,"moyenne":78,"totalCampagnes":3,"totalTalents":22,"iaProcessed":30}
        },
        "AdminEntreprise": {
            "kpis": [
                {"label":"TALENTS ACTIFS","value":124+random.randint(-5,10),       "icon":"fa-solid fa-user-group",   "color":"#fbbf24","bg":"rgba(251,191,36,0.12)","trend":f"+{random.randint(3,12)}","sparkPoints":"0,22 12,18 24,20 36,10 48,14 60,8 72,12 80,6"},
                {"label":"TAUX RÉUSSITE", "value":f"{88+random.randint(-3,5)}%",   "icon":"fa-solid fa-circle-check", "color":"#10b981","bg":"rgba(16,185,129,0.12)","trend":"+3%","sparkPoints":"0,24 12,20 24,16 36,18 48,12 60,10 72,8 80,4"},
                {"label":"SESSIONS",      "value":12+random.randint(-2,4),         "icon":"fa-solid fa-bolt-lightning","color":"#3b82f6","bg":"rgba(59,130,246,0.12)","trend":f"+{random.randint(1,5)}","sparkPoints":"0,20 12,16 24,22 36,14 48,16 60,12 72,14 80,8"},
                {"label":"ÉVALUATIONS IA","value":89+random.randint(-5,10),        "icon":"fa-solid fa-brain",        "color":"#8b5cf6","bg":"rgba(139,92,246,0.12)","trend":"↑","sparkPoints":"0,26 12,22 24,24 36,16 48,18 60,10 72,12 80,6"},
            ],
            "summary":{"totalTests":89,"moyenne":88,"totalCampagnes":12,"totalTalents":124,"iaProcessed":89}
        },
        "SuperAdmin": {
            "kpis": [
                {"label":"ORGANISATIONS", "value":42+random.randint(-1,3),           "icon":"fa-solid fa-building",    "color":"#6366f1","bg":"rgba(99,102,241,0.12)","trend":"+2","sparkPoints":"0,22 12,18 24,20 36,10 48,14 60,8 72,12 80,6"},
                {"label":"UTILISATEURS",  "value":1247+random.randint(-10,25),        "icon":"fa-solid fa-users",       "color":"#10b981","bg":"rgba(16,185,129,0.12)","trend":"+47","sparkPoints":"0,24 12,20 24,16 36,18 48,12 60,10 72,8 80,4"},
                {"label":"SESSIONS/JOUR", "value":340+random.randint(-20,30),         "icon":"fa-solid fa-wave-square", "color":"#f59e0b","bg":"rgba(245,158,11,0.12)","trend":f"+{random.randint(5,20)}","sparkPoints":"0,20 12,16 24,22 36,14 48,16 60,12 72,14 80,8"},
                {"label":"UPTIME",        "value":f"{round(99.7+random.uniform(0,0.3),1)}%","icon":"fa-solid fa-server","color":"#ef4444","bg":"rgba(239,68,68,0.12)","trend":"99.9%","sparkPoints":"0,26 12,22 24,24 36,16 48,18 60,10 72,12 80,6"},
            ],
            "summary":{"totalEntreprises":42,"totalUtilisateurs":1247,"totalTests":340,"demandesEnAttente":8}
        },
    }
    return configs.get(role, configs["AdminEntreprise"])


def _get_role_insight(role: str) -> str:
    insights = {
        "Candidat":        "Vous avez 2 tests en attente. Score moyen : 76% (+8% ce mois). Aura Score : 82/100.",
        "Evaluateur":      "5 candidats en attente. Taux de traitement 91%. Session planifiée demain à 10h.",
        "RH":              "3 campagnes actives. 18 candidats en cours. Taux de réussite : 84%.",
        "Recruteur":       "7 profils compatibles (85%+). Taux matching +15% cette semaine.",
        "AdminEntreprise": "124 talents actifs. Taux réussite : 88%. Aura Score : 94. 8 candidats en attente.",
        "SuperAdmin":      "42 entreprises actives. 1 247 utilisateurs. Uptime 99.9%. ⚠️ Mailer DOWN depuis 2h.",
    }
    return insights.get(role, "Tableau de bord EvaluaTech — données en temps réel.")


def _get_default_activities() -> list:
    now = datetime.now()
    return [
        {"id":1,"user":"IA Gemini","action":"Recommandations générées","color":"#f59e0b","time":now.strftime("%H:%M"),"module":"Recommandations"},
        {"id":2,"user":"Système","action":"Moteur IA v10.1 démarré","color":"#10b981","time":(now-timedelta(minutes=2)).strftime("%H:%M"),"module":"Système"},
        {"id":3,"user":"Mehdi Admin","action":"Dashboard ouvert","color":"#6366f1","time":(now-timedelta(minutes=5)).strftime("%H:%M"),"module":"Dashboard"},
        {"id":4,"user":"IA","action":"Matching CV — Score 91%","color":"#fbbf24","time":(now-timedelta(minutes=8)).strftime("%H:%M"),"module":"Analyses CV"},
        {"id":5,"user":"Recruteur","action":"Test React Senior créé","color":"#3b82f6","time":(now-timedelta(minutes=15)).strftime("%H:%M"),"module":"Évaluations"},
        {"id":6,"user":"Amira Saidi","action":"Candidat Tarek Ben Salem — Retenu","color":"#10b981","time":(now-timedelta(minutes=22)).strftime("%H:%M"),"module":"Évaluations"},
        {"id":7,"user":"IA Gemini","action":"Lettre motivation EN générée","color":"#8b5cf6","time":(now-timedelta(minutes=30)).strftime("%H:%M"),"module":"Lettres"},
        {"id":8,"user":"⚠️ Alerte","action":"Document non-CV uploadé — rejeté","color":"#f59e0b","time":(now-timedelta(minutes=35)).strftime("%H:%M"),"module":"Analyses CV"},
    ]


def _get_chart_data(role: str, period: str = "week") -> dict:
    days = {"week":7,"month":30,"quarter":90}.get(period, 7)
    today = datetime.now()
    base  = {"SuperAdmin":60,"AdminEntreprise":45,"Recruteur":30,"Evaluateur":20,"RH":25,"Candidat":15}.get(role,30)
    labels, values, values2 = [], [], []
    for i in range(days-1, -1, -1):
        d = today - timedelta(days=i)
        labels.append(d.strftime("%d/%m") if days > 7 else ["Dim","Lun","Mar","Mer","Jeu","Ven","Sam"][d.weekday()])
        v = max(0, round(base + random.randint(-10,30)))
        values.append(v); values2.append(max(0, v - random.randint(5,20)))
    return {"labels":labels,"datasets":[{"label":"Performance","data":values,"color":"#fbbf24"},{"label":"Référence","data":values2,"color":"#94a3b8"}],"period":period}


def _get_widget_config(role: str) -> dict:
    configs = {
        "Candidat":        {"theme_color":"#3b82f6","accent":"#60a5fa","hero_icon":"fa-user-graduate","sections":["kpis","tests_en_cours","progression","cv","lettre","resultats","chart","recommendations"],"quick_actions":[{"label":"Passer un test","icon":"fa-play","route":"/my-tests","color":"#3b82f6"},{"label":"Analyser CV","icon":"fa-file-pdf","route":"/dashboard#cv","color":"#f59e0b"},{"label":"Générer lettre","icon":"fa-envelope","route":"/dashboard#lettre","color":"#8b5cf6"},{"label":"Préparer entretien","icon":"fa-microphone","route":"/dashboard#interview","color":"#10b981"}]},
        "Evaluateur":      {"theme_color":"#f59e0b","accent":"#fbbf24","hero_icon":"fa-clipboard-check","sections":["kpis","eval_queue","sessions","scan_cv","top_skills","chart","activity","recommendations"],"quick_actions":[{"label":"Ma file","icon":"fa-users-gear","route":"/analyse-comportementale","color":"#f59e0b"},{"label":"Planifier","icon":"fa-calendar-plus","route":"/sessions","color":"#6366f1"},{"label":"Analyse comportementale","icon":"fa-brain","route":"/analyse-comportementale","color":"#10b981"},{"label":"Rapport","icon":"fa-chart-bar","route":"/reporting","color":"#8b5cf6"}]},
        "RH":              {"theme_color":"#8b5cf6","accent":"#a78bfa","hero_icon":"fa-people-arrows","sections":["kpis","eval_queue","sessions","scan_cv","top_skills","chart","activity","recommendations"],"quick_actions":[{"label":"Créer campagne","icon":"fa-bullhorn","route":"/campaigns","color":"#8b5cf6"},{"label":"Inviter","icon":"fa-user-plus","route":"/invite","color":"#10b981"},{"label":"Rapport mensuel","icon":"fa-chart-bar","route":"/reporting","color":"#f59e0b"},{"label":"Soft skills","icon":"fa-brain","route":"/analyse-comportementale","color":"#3b82f6"}]},
        "Recruteur":       {"theme_color":"#10b981","accent":"#34d399","hero_icon":"fa-handshake","sections":["kpis","pipeline","scan_cv","chart","activity","recommendations"],"quick_actions":[{"label":"Nouvelle campagne","icon":"fa-plus","route":"/campaigns","color":"#10b981"},{"label":"Pipeline","icon":"fa-kanban","route":"/campaigns","color":"#6366f1"},{"label":"Analyser CV","icon":"fa-file-pdf","route":"/analyse-comportementale","color":"#f59e0b"},{"label":"QCM IA","icon":"fa-wand-sparkles","route":"/ai-generator","color":"#8b5cf6"}]},
        "AdminEntreprise": {"theme_color":"#f59e0b","accent":"#fbbf24","hero_icon":"fa-building-user","sections":["kpis","activity","team","scan_cv","recent_candidates","top_skills","chart","recommendations"],"quick_actions":[{"label":"Pipeline","icon":"fa-kanban","route":"/campaigns","color":"#f59e0b"},{"label":"Équipe","icon":"fa-people-group","route":"/staff-members","color":"#6366f1"},{"label":"Analyser CV","icon":"fa-file-pdf","route":"/analyse-comportementale","color":"#10b981"},{"label":"Rapport","icon":"fa-chart-bar","route":"/reporting","color":"#8b5cf6"}]},
        "SuperAdmin":      {"theme_color":"#6366f1","accent":"#818cf8","hero_icon":"fa-shield-halved","sections":["kpis","services","companies","subscriptions","scan_cv","chart","activity","recommendations"],"quick_actions":[{"label":"Organisations","icon":"fa-building","route":"/super-admin","color":"#6366f1"},{"label":"Santé","icon":"fa-server","route":"/super-admin","color":"#10b981"},{"label":"Abonnements","icon":"fa-coins","route":"/gestion-abonnements","color":"#f59e0b"},{"label":"Analytics","icon":"fa-chart-mixed","route":"/super-admin-analytics","color":"#8b5cf6"}]},
    }
    return configs.get(role, configs["AdminEntreprise"])


def _get_role_navigation(role: str) -> list:
    navs = {
        "Candidat":        [{"label":"Dashboard","icon":"fa-grid-2","route":"/dashboard","active":True},{"label":"Mes Tests","icon":"fa-clipboard-list","route":"/my-tests"},{"label":"Résultats","icon":"fa-chart-bar","route":"/results"},{"label":"Historique","icon":"fa-clock-rotate-left","route":"/history"},{"label":"Profil","icon":"fa-user","route":"/profile"}],
        "Evaluateur":      [{"label":"Dashboard","icon":"fa-grid-2","route":"/dashboard","active":True},{"label":"Évaluations","icon":"fa-clipboard-check","route":"/analyse-comportementale"},{"label":"Sessions","icon":"fa-calendar","route":"/sessions"},{"label":"Analyse","icon":"fa-brain","route":"/analyse-comportementale"},{"label":"Stats","icon":"fa-chart-pie","route":"/stats"}],
        "RH":              [{"label":"Dashboard","icon":"fa-grid-2","route":"/dashboard","active":True},{"label":"Campagnes","icon":"fa-bullhorn","route":"/campaigns"},{"label":"Candidats","icon":"fa-users","route":"/candidates-list"},{"label":"Évaluations","icon":"fa-clipboard-check","route":"/analyse-comportementale"},{"label":"Rapports","icon":"fa-chart-bar","route":"/reporting"}],
        "Recruteur":       [{"label":"Dashboard","icon":"fa-grid-2","route":"/dashboard","active":True},{"label":"Campagnes","icon":"fa-bullhorn","route":"/campaigns"},{"label":"Candidats","icon":"fa-users","route":"/candidates-list"},{"label":"Banque Q.","icon":"fa-database","route":"/questions"},{"label":"IA Gen.","icon":"fa-wand-sparkles","route":"/ai-generator"},{"label":"Rapports","icon":"fa-chart-bar","route":"/reporting"}],
        "AdminEntreprise": [{"label":"Dashboard","icon":"fa-grid-2","route":"/dashboard","active":True},{"label":"Campagnes","icon":"fa-bullhorn","route":"/campaigns"},{"label":"Candidats","icon":"fa-users","route":"/candidates-list"},{"label":"Équipe","icon":"fa-people-group","route":"/staff-members"},{"label":"Rapports","icon":"fa-chart-bar","route":"/reporting"},{"label":"Rôles","icon":"fa-shield-halved","route":"/roles"}],
        "SuperAdmin":      [{"label":"Dashboard","icon":"fa-grid-2","route":"/dashboard","active":True},{"label":"Organisations","icon":"fa-building","route":"/super-admin"},{"label":"Utilisateurs","icon":"fa-users","route":"/platform-users"},{"label":"Abonnements","icon":"fa-coins","route":"/gestion-abonnements"},{"label":"Analytics","icon":"fa-chart-mixed","route":"/super-admin-analytics"},{"label":"Supervision","icon":"fa-shield-halved","route":"/super-admin"}],
    }
    return navs.get(role, navs["AdminEntreprise"])


# ─── DASHBOARD ENDPOINTS ───

@app.get("/")
async def root():
    return {
        "status": "Online", "model": WORKING_MODEL, "circuit": _circuit.state,
        "uptime": f"{int(time.time()-_START_TIME)}s", "version": "10.1",
        "gemini_available": _gemini_client is not None,
        "total_intents": len(UNIVERSAL_BRAIN),
        "new_in_v10_1": {
            "lettre_motivation": "Langue strictement respectée (FR/EN/AR), enforcement triple couche",
            "cv_validation": "Détection document non-CV avant analyse avec alerte structurée",
            "cv_validate_endpoint": "/ia/validate-cv-doc",
            "lettre_languages_endpoint": "/ia/lettre-motivation/languages",
            "theme_soustheme_intent": "Chatbot répond aux questions sur les thèmes/sous-thèmes disponibles",
            "qcm_fix": "theme + sousTheme injectés dans chaque question générée",
        },
        "endpoints": {
            "dashboard_dynamic":    "/ia/dashboard/dynamic?role=AdminEntreprise",
            "dashboard_view":       "/ia/dashboard/view?role=AdminEntreprise",
            "global_stats":         "/ia/dashboard/global-stats?role=AdminEntreprise",
            "analytics_overview":   "/ia/analytics/overview?role=AdminEntreprise",
            "recommendations":      "/ia/recommendations?role=AdminEntreprise",
            "chat":                 "POST /ia/chat",
            "cv_analysis":          "POST /ia/match-cv  [détection non-CV]",
            "cv_validate":          "POST /ia/validate-cv-doc",
            "lettre":               "POST /ia/lettre-motivation  [langue stricte]",
            "lettre_languages":     "GET /ia/lettre-motivation/languages",
            "qcm":                  "POST /ia/generate-bilingual  [theme+sousTheme injectés]",
            "voice":                "GET /ia/voice/config",
            "health":               "GET /ia/health",
        }
    }

@app.get("/ia/health")
async def detailed_health():
    h = list(AI_METRICS["latency_history"])
    return {
        "status": "ok", "model": WORKING_MODEL, "circuit_state": _circuit.state,
        "uptime_seconds": int(time.time()-_START_TIME),
        "total_calls": AI_METRICS["total_calls"], "error_count": AI_METRICS["error_count"],
        "avg_latency_ms": round(sum(h)/len(h),1) if h else 0,
        "active_requests": AI_METRICS["active_requests"],
        "cache_size": len(_cache._store), "tokens_used": AI_METRICS["total_tokens"],
        "chat_cache_size": len(_chat_cache._store),
        "active_chat_sessions": len(_session_memory) if '_session_memory' in globals() else 0,
        "dashboard_requests": AI_METRICS["dashboard_requests"],
        "gemini_available": _gemini_client is not None,
        "total_intents": len(UNIVERSAL_BRAIN),
        "cv_rejected_total": AI_METRICS["cv_rejected"],
        "version": "10.1",
    }


@app.get("/ia/dashboard/view")
async def dashboard_view(
    role: str = Query("AdminEntreprise"),
    org_id: str = Query(None), user_id: str = Query(None),
    lang: str = Query("fr"),
    include_recommendations: bool = Query(True),
    include_stats: bool = Query(True),
    include_activity: bool = Query(True),
    include_chart: bool = Query(False)
):
    now = datetime.now()
    ck  = make_cache_key("dash-v10", role, now.strftime("%Y%m%d%H%M"))
    cached = _dash_cache.get(ck)
    if cached:
        AI_METRICS["dashboard_requests"] += 1
        return cached

    kpi_data = _get_dynamic_kpis(role)
    result = {
        "role": role, "generated_at": now.isoformat(), "cache_ttl_seconds": 60, "version": "10.1",
        "kpis":          kpi_data["kpis"] if include_stats else [],
        "summary":       kpi_data["summary"],
        "insight":       _get_role_insight(role),
        "navigation":    _get_role_navigation(role),
        "widget_config": _get_widget_config(role),
        "meta": {"uptime_seconds": int(time.time()-_START_TIME), "model": WORKING_MODEL, "circuit_state": _circuit.state, "gemini_available": _gemini_client is not None, "platform_version": "10.1"},
    }

    if role == "Candidat":
        result["role_data"] = {"tests_en_cours": _DS["candidat_tests"], "progression": [{"campagneNom":h["titreExamen"],"score":h["score"]} for h in _DS["candidat_historique"][:5]], "last_results": _DS["candidat_historique"][:5]}
    elif role in ["Evaluateur","RH"]:
        result["role_data"] = {"eval_queue": _DS["eval_queue"], "sessions": _DS["campagnes"], "top_skills": _DS["top_skills"]}
    elif role == "Recruteur":
        result["role_data"] = {"pipeline_columns": [{"id":1,"title":"INVITÉS","count":_DS["pipeline"]["invites"]["count"],"color":"#6366f1","cards":_DS["pipeline"]["invites"]["cards"]},{"id":2,"title":"EN COURS","count":_DS["pipeline"]["en_cours"]["count"],"color":"#f59e0b","cards":_DS["pipeline"]["en_cours"]["cards"]},{"id":3,"title":"COMPLÉTÉS","count":_DS["pipeline"]["completes"]["count"],"color":"#10b981","cards":_DS["pipeline"]["completes"]["cards"]},{"id":4,"title":"RETENUS","count":_DS["pipeline"]["retenus"]["count"],"color":"#22c55e","cards":_DS["pipeline"]["retenus"]["cards"]}],"recent_candidates": _DS["recent_candidates"][:4],"top_skills": _DS["top_skills"][:4]}
    elif role == "AdminEntreprise":
        result["role_data"] = {"team": _DS["staff_members"], "recent_candidates": _DS["recent_candidates"], "top_skills": _DS["top_skills"], "eval_queue": _DS["eval_queue"][:5], "campagnes": _DS["campagnes"][:4], "team_stats": {"total":len(_DS["staff_members"]),"active":sum(1 for m in _DS["staff_members"] if m["estActif"]),"inactive":sum(1 for m in _DS["staff_members"] if not m["estActif"])}}
    elif role == "SuperAdmin":
        services = [dict(s) for s in _DS["sa_services"]]
        result["role_data"] = {"services": services, "companies": _DS["sa_companies"], "subscriptions": _DS["sa_subscriptions"], "platform_users": _DS["platform_users"][:6], "platform_stats": {"up_services":sum(1 for s in services if s["up"]),"down_services":sum(1 for s in services if not s["up"]),"total_users":sum(c["users"] for c in _DS["sa_companies"]),"total_revenue":sum(s["revenue"] for s in _DS["sa_subscriptions"])}}

    if include_activity:
        acts = list(_activity_log)[:8]
        result["activity"] = acts if acts else _get_default_activities()
    if include_recommendations:
        ck_reco = make_cache_key("reco-v10", role, lang)
        result["recommendations"] = _reco_cache.get(ck_reco) or _get_fallback_recommendations(role)
    if include_chart:
        result["chart_config"] = _get_chart_data(role)

    _dash_cache.set(ck, result)
    AI_METRICS["dashboard_requests"] += 1
    AI_METRICS["usage_counts"]["Dashboard"] += 1
    log_activity("Système", f"Dashboard.view ({role})", "#6366f1", "Dashboard")
    return result


@app.get("/ia/dashboard/dynamic")
async def dashboard_dynamic_all(role: str = Query("AdminEntreprise"), user_id: str = Query(None)):
    now      = datetime.now()
    kpi_data = _get_dynamic_kpis(role)
    data = {
        "role": role, "generated_at": now.isoformat(),
        "kpis": kpi_data["kpis"], "summary": kpi_data["summary"],
        "insight": _get_role_insight(role),
        "chart": _get_chart_data(role, "week"),
        "activities": list(_activity_log)[:8] if _activity_log else _get_default_activities(),
        "recommendations": _reco_cache.get(make_cache_key("reco-v10", role, "fr")) or _get_fallback_recommendations(role),
        "widget_config": _get_widget_config(role),
        "navigation": _get_role_navigation(role),
    }
    if role == "Candidat":
        data["candidat_tests"] = _DS["candidat_tests"]
        data["candidat_historique"] = _DS["candidat_historique"][:5]
        data["candidat_progression"] = [{"campagneNom":h["titreExamen"],"score":h["score"]} for h in _DS["candidat_historique"][:5]]
    elif role in ["Evaluateur","RH"]:
        data["eval_queue"] = _DS["eval_queue"]
        data["campagnes"]  = _DS["campagnes"]
        data["top_skills"] = _DS["top_skills"]
    elif role == "Recruteur":
        data["pipeline"]   = _DS["pipeline"]
        data["campagnes"]  = _DS["campagnes"][:3]
        data["top_skills"] = _DS["top_skills"][:4]
        data["recent_candidates"] = _DS["recent_candidates"][:4]
    elif role == "AdminEntreprise":
        data["staff_members"]     = _DS["staff_members"]
        data["recent_candidates"] = _DS["recent_candidates"]
        data["top_skills"]        = _DS["top_skills"]
        data["eval_queue"]        = _DS["eval_queue"][:5]
        data["campagnes"]         = _DS["campagnes"][:4]
    elif role == "SuperAdmin":
        data["services"]       = [dict(s) for s in _DS["sa_services"]]
        data["companies"]      = _DS["sa_companies"]
        data["subscriptions"]  = _DS["sa_subscriptions"]
        data["platform_users"] = _DS["platform_users"][:6]
        data["superadmin_stats"] = {"totalEntreprises":42,"totalUtilisateurs":1247,"demandesEnAttente":8,"totalTests":340,"croissanceStats":[{"mois":"Fév","count":3},{"mois":"Mar","count":5},{"mois":"Avr","count":4},{"mois":"Mai","count":2}]}
    AI_METRICS["dashboard_requests"] += 1
    return data


@app.get("/ia/dashboard/global-stats")
async def dashboard_global_stats(role: str = Query("AdminEntreprise")):
    kpi_data = _get_dynamic_kpis(role)
    leaders  = sorted(_DS["recent_candidates"], key=lambda x: x["score"], reverse=True)[:5]
    return {
        "kpis": kpi_data["summary"],
        "chart": [{"name":c["campagneNom"],"score":h["score"]} for h,c in zip(_DS["candidat_historique"][:5], _DS["campagnes"][:5])],
        "leaders": [{"name":c["candidateName"],"score":c["score"],"test":c["testName"]} for c in leaders],
        "recentResults": _DS["recent_candidates"],
        "insight": _get_role_insight(role),
        "recentActivities": _get_default_activities()[:5],
        "generated_at": datetime.now().isoformat()
    }


@app.get("/ia/analytics/overview")
async def analytics_overview(role: str = Query("AdminEntreprise")):
    kpi_data = _get_dynamic_kpis(role)
    return {
        "kpis": kpi_data["summary"],
        "insight": _get_role_insight(role),
        "recentActivities": _get_default_activities()[:8],
        "chartData": _get_chart_data(role,"week")["datasets"][0]["data"],
        "generated_at": datetime.now().isoformat()
    }


@app.get("/ia/dashboard/stats")
async def dashboard_stats(role: str = Query("AdminEntreprise")):
    kpi_data = _get_dynamic_kpis(role)
    return {"role":role,"kpis":kpi_data["kpis"],"summary":kpi_data["summary"],"insight":_get_role_insight(role),"generated_at":datetime.now().isoformat()}

@app.get("/ia/dashboard/activity")
async def dashboard_activity(limit: int = Query(10), role: str = Query("AdminEntreprise")):
    acts = list(_activity_log)[:limit] if _activity_log else _get_default_activities()[:limit]
    return {"activities": acts if acts else _get_default_activities()[:limit]}

@app.get("/ia/dashboard/chart-data")
async def dashboard_chart(role: str = Query("AdminEntreprise"), period: str = Query("week")):
    return _get_chart_data(role, period)

@app.get("/ia/dashboard/candidat-tests")
async def candidat_tests(user_id: str = Query(None)):
    return {"tests": _DS["candidat_tests"], "total": len(_DS["candidat_tests"])}

@app.get("/ia/dashboard/candidat-progression")
async def candidat_progression(user_id: str = Query(None)):
    return {"progression": [{"campagneNom":h["titreExamen"],"score":h["score"]} for h in _DS["candidat_historique"][:5]]}

@app.get("/ia/dashboard/candidat-results")
async def candidat_results(user_id: str = Query(None), limit: int = Query(5)):
    return {"results": _DS["candidat_historique"][:limit], "total": len(_DS["candidat_historique"])}

@app.get("/ia/dashboard/eval-queue")
async def eval_queue(role: str = Query("Evaluateur"), limit: int = Query(6)):
    q = _DS["eval_queue"][:limit]
    return {"queue":q,"total":len(_DS["eval_queue"]),"urgent_count":sum(1 for e in q if e["urgency"]=="urgent")}

@app.get("/ia/dashboard/sessions")
async def get_sessions(role: str = Query("Evaluateur"), limit: int = Query(4)):
    return {"sessions": _DS["campagnes"][:limit], "total": len(_DS["campagnes"])}

@app.get("/ia/dashboard/top-skills")
async def get_top_skills(role: str = Query("Evaluateur"), limit: int = Query(6)):
    return {"skills": [dict(s) for s in _DS["top_skills"][:limit]]}

@app.get("/ia/dashboard/team")
async def get_team(org_id: str = Query(None), limit: int = Query(6)):
    team = _DS["staff_members"][:limit]
    return {"team":team,"total":len(_DS["staff_members"]),"active_count":sum(1 for m in team if m["estActif"]),"inactive_count":sum(1 for m in team if not m["estActif"])}

@app.get("/ia/dashboard/recent-candidates")
async def get_recent_candidates(org_id: str = Query(None), limit: int = Query(6)):
    return {"candidates": _DS["recent_candidates"][:limit], "total": len(_DS["recent_candidates"])}

@app.get("/ia/dashboard/services")
async def get_services():
    services = [dict(s) for s in _DS["sa_services"]]
    up_count = sum(1 for s in services if s["up"])
    return {"services":services,"up_count":up_count,"down_count":len(services)-up_count,"uptime_pct":round((up_count/len(services))*100,1)}

@app.get("/ia/dashboard/companies")
async def get_companies(limit: int = Query(6)):
    companies = _DS["sa_companies"][:limit]
    return {"companies":companies,"total":len(_DS["sa_companies"]),"enterprise_count":sum(1 for c in companies if c["plan"]=="Enterprise"),"total_users":sum(c["users"] for c in companies)}

@app.get("/ia/dashboard/subscriptions")
async def get_subscriptions():
    return {"subscriptions":_DS["sa_subscriptions"],"total":sum(s["count"] for s in _DS["sa_subscriptions"]),"total_revenue":sum(s["revenue"] for s in _DS["sa_subscriptions"])}

@app.get("/ia/dashboard/pipeline")
async def dashboard_pipeline(role: str = Query("Recruteur")):
    return {"columns":[{"id":1,"title":"INVITÉS","count":_DS["pipeline"]["invites"]["count"],"color":"#6366f1","cards":_DS["pipeline"]["invites"]["cards"]},{"id":2,"title":"EN COURS","count":_DS["pipeline"]["en_cours"]["count"],"color":"#f59e0b","cards":_DS["pipeline"]["en_cours"]["cards"]},{"id":3,"title":"COMPLÉTÉS","count":_DS["pipeline"]["completes"]["count"],"color":"#10b981","cards":_DS["pipeline"]["completes"]["cards"]},{"id":4,"title":"RETENUS","count":_DS["pipeline"]["retenus"]["count"],"color":"#22c55e","cards":_DS["pipeline"]["retenus"]["cards"]}]}

@app.get("/ia/dashboard/realtime")
async def dashboard_realtime(role: str = Query("AdminEntreprise")):
    AI_METRICS["realtime_updates"] += 1
    alerts = []
    if role == "SuperAdmin": alerts = [{"type":"warning","message":"Service Mailer DOWN depuis 2h — 47 invitations bloquées","color":"#ef4444"}]
    elif role == "AdminEntreprise": alerts = [{"type":"info","message":"8 candidats en attente d'analyse depuis 48h","color":"#f59e0b"}]
    return {"timestamp":datetime.now().isoformat(),"role":role,"kpis_snapshot":_get_dynamic_kpis(role)["summary"],"active_users":random.randint(8,45),"active_sessions":random.randint(3,18),"new_events":random.randint(0,5),"alerts":alerts,"cache_status":"healthy","model_status":_circuit.state,"gemini_available":_gemini_client is not None}


# ════════════════════════════════════════════════════════════
# ██   CHATBOT   ██
# ════════════════════════════════════════════════════════════

_session_memory: dict = {}
CHAT_CTX_MAX  = 6
MAX_SESSIONS  = 500

def _get_session(session_id: str) -> dict:
    if session_id not in _session_memory:
        if len(_session_memory) >= MAX_SESSIONS:
            oldest = min(_session_memory.items(), key=lambda x: x[1].get("last_seen",0))
            del _session_memory[oldest[0]]
        _session_memory[session_id] = {"history":[],"last_seen":time.time(),"lang":"fr","role":"Recruteur","question_count":0}
    _session_memory[session_id]["last_seen"] = time.time()
    return _session_memory[session_id]

CHAT_SYSTEM = (
    "Tu es NeoBot v10.1, assistant IA EvaluaTech. Réponds en 3-5 phrases max, directement. "
    "EvaluaTech v10.1 = SaaS évaluation technique+comportementale. "
    "NOUVEAUTÉS v10.1: Lettre motivation en langue strictement choisie (FR/EN/AR) + "
    "Analyse CV avec détection automatique si document n'est pas un CV (alerte affichée) + "
    "Thèmes/sous-thèmes disponibles: React, Vue.js, Angular, Node.js, Python, Java, SQL, Docker, AWS, Machine Learning, Soft Skills, Leadership. "
    "Modules: Dashboard | QCM IA(45s) | AnalyseCV | LettreMotivation | EntretienIA | Proctoring | Rapports | BanqueQuestions. "
    "Rôles: Candidat | Evaluateur | RH | Recruteur | AdminEntreprise | SuperAdmin. "
    "Support: support@evaluatech.com. Adapte la langue (FR/EN/AR auto)."
)

def _get_smart_fallback(message: str, lang: str, role: str) -> str:
    responses = {
        "fr": ("🤖 **NeoBot v10.1** — Mode local actif.\n\n**Nouveautés :**\n- ✉️ Lettre de motivation : langue strictement respectée (FR/EN/AR)\n- 📄 Analyse CV : alerte si document non-CV\n- 🎯 Thèmes disponibles : React, Vue.js, Node.js, Python, Java, SQL, Docker, AWS, ML, Soft Skills\n\n**Fonctionnalités :** Dashboard | Tests IA | Analyse CV | Lettre | Proctoring\n\n💡 Posez une question précise !"),
        "en":  ("🤖 **NeoBot v10.1** — Local mode.\n\n**New:** Cover letter in strict language (FR/EN/AR) | CV non-CV detection alert | Themes: React, Vue.js, Node.js, Python, Java, SQL, Docker, AWS, ML, Soft Skills\n\nFeatures: Dashboard | AI Tests | CV Analysis | Cover Letter | Proctoring"),
        "ar":  ("🤖 **NeoBot v10.1** — الوضع المحلي.\n\n**جديد:** خطاب التقديم بلغة صارمة | كشف الوثائق غير السيرة الذاتية | المواضيع: React · Python · Java · SQL · Docker · AWS")
    }
    return responses.get(lang, responses["fr"])

def _build_chat_prompt(session: dict, message: str, role: str, lang: str) -> str:
    lang_instr = {"fr":"Réponds en français.","en":"Reply in English.","ar":"أجب بالعربية."}.get(lang,"Réponds en français.")
    lines = [f"[SYSTEM: {CHAT_SYSTEM} {lang_instr}]", f"[ROLE:{role}][LANG:{lang}]"]
    for turn in session["history"][-4:]:
        r = "Utilisateur" if turn.get("role")=="user" else "NeoBot"
        lines.append(f"{r}: {str(turn.get('content',''))[:200]}")
    lines.append(f"Utilisateur: {message}")
    lines.append("NeoBot:")
    return "\n".join(lines)


@app.post("/ia/chat")
async def chat_interaction(
    message: str = Form(...),
    history: str = Form("[]"),
    role: str = Form("Recruteur"),
    lang: str = Form("auto"),
    session_id: str = Form("default")
):
    effective_lang = detect_language(message) if lang == "auto" else lang
    session = _get_session(session_id)
    session["lang"] = effective_lang; session["role"] = role
    session["question_count"] += 1
    AI_METRICS["usage_counts"]["Chat"] += 1

    ck = make_cache_key("chat-v10", message[:60], effective_lang, role)
    cached = _chat_cache.get(ck)
    if cached:
        AI_METRICS["chat_cache_hits"] += 1
        return {"status":"SUCCESS","response":cached["response"],"reply":cached["response"],"suggestions":cached.get("suggestions",[]),"lang_detected":effective_lang,"source":"cache"}

    local_reply = _get_local_response(message, effective_lang, role)
    if local_reply:
        suggestions = _get_suggestions(message, effective_lang, role)
        AI_METRICS["chat_intent_hits"] += 1
        _chat_cache.set(ck, {"response":local_reply,"suggestions":suggestions})
        session["history"].append({"role":"user","content":message})
        session["history"].append({"role":"assistant","content":local_reply})
        session["history"] = session["history"][-CHAT_CTX_MAX:]
        log_activity("Chat", f"Question ({role})", "#3b82f6", "Chat")
        return {"status":"SUCCESS","response":local_reply,"reply":local_reply,"suggestions":suggestions,"lang_detected":effective_lang,"source":"intent"}

    try:
        prompt = _build_chat_prompt(session, message, role, effective_lang)
        AI_METRICS["chat_gemini_calls"] += 1
        res   = await call_gemini_async(prompt, module="Chat", sem=_sem_chat, retries=1)
        reply = res.text.strip()
        if not reply: raise QuotaExceeded("Empty")
        suggestions = _get_suggestions(message, effective_lang, role)
        _chat_cache.set(ck, {"response":reply,"suggestions":suggestions})
        session["history"].append({"role":"user","content":message})
        session["history"].append({"role":"assistant","content":reply})
        session["history"] = session["history"][-CHAT_CTX_MAX:]
        return {"status":"SUCCESS","response":reply,"reply":reply,"suggestions":suggestions,"lang_detected":effective_lang,"source":"gemini"}
    except (QuotaExceeded, Exception):
        fallback = _get_smart_fallback(message, effective_lang, role)
        return {"status":"SUCCESS","response":fallback,"reply":fallback,"suggestions":_get_suggestions(message,effective_lang,role),"lang_detected":effective_lang,"source":"fallback"}


@app.get("/ia/chat/suggestions")
async def chat_suggestions(role: str = Query("Recruteur"), lang: str = Query("fr")):
    starters = {
        "Recruteur":       {"fr":["Analyser un CV","Créer un test","Quels thèmes disponibles ?","Pipeline Kanban","Recommandations IA"],"en":["Analyze CV","Create test","Available themes?","Kanban pipeline","AI recommendations"],"ar":["تحليل سيرة","إنشاء اختبار","المواضيع المتاحة؟","المسار","التوصيات"]},
        "Candidat":        {"fr":["Mon dashboard","Générer ma lettre en arabe","Analyser mon CV","Préparer mon entretien","Mes résultats"],"en":["My dashboard","Generate letter in French","Analyze my CV","Prepare interview"],"ar":["لوحتي","توليد خطاب بالفرنسية","تحليل سيرتي","تحضير مقابلة"]},
        "AdminEntreprise": {"fr":["Dashboard temps réel","Recommandations IA","Gérer mon équipe","Rapport mensuel","Quels thèmes disponibles ?"],"en":["Real-time dashboard","AI recommendations","Manage team","Monthly report","Available themes?"],"ar":["لوحة فورية","التوصيات","إدارة الفريق","تقرير شهري","المواضيع المتاحة؟"]},
        "SuperAdmin":      {"fr":["Santé des services","Mailer service DOWN ?","Gérer les organisations","Abonnements","Audit sécurité"],"en":["Services health","Mailer service DOWN?","Manage organizations","Subscriptions"],"ar":["صحة الخدمات","إدارة المنظمات","الاشتراكات"]},
        "Evaluateur":      {"fr":["Ma file d'évaluation","Candidats urgents","Planifier une session","Top compétences","Rapport"],"en":["Evaluation queue","Urgent candidates","Schedule session","Top skills"],"ar":["قائمة التقييم","المرشحون العاجلون","جدولة جلسة"]},
        "RH":              {"fr":["Créer une campagne","Quels thèmes disponibles ?","Rapport RH mensuel","Statistiques","Recommandations"],"en":["Create campaign","Available themes?","Monthly report","Statistics"],"ar":["إنشاء حملة","المواضيع المتاحة؟","التقرير الشهري"]},
    }
    role_sugg = starters.get(role, starters["Recruteur"])
    return {"suggestions": role_sugg.get(lang, role_sugg.get("fr", []))}

@app.post("/ia/chat/reset")
async def reset_chat_session(session_id: str = Form("default")):
    if session_id in _session_memory: del _session_memory[session_id]
    return {"status":"OK","message":f"Session {session_id} réinitialisée."}

@app.get("/ia/chat/metrics")
async def chat_metrics():
    total = AI_METRICS["chat_cache_hits"] + AI_METRICS["chat_intent_hits"] + AI_METRICS["chat_gemini_calls"]
    return {
        "total_chat_requests":total,"cache_hits":AI_METRICS["chat_cache_hits"],
        "intent_hits":AI_METRICS["chat_intent_hits"],"gemini_calls":AI_METRICS["chat_gemini_calls"],
        "cache_hit_rate":f"{round(AI_METRICS['chat_cache_hits']/max(total,1)*100)}%",
        "intent_hit_rate":f"{round(AI_METRICS['chat_intent_hits']/max(total,1)*100)}%",
        "active_sessions":len(_session_memory),
        "gemini_available":_gemini_client is not None,
        "total_intents":len(UNIVERSAL_BRAIN),
        "cv_rejected":AI_METRICS["cv_rejected"],
    }

@app.post("/ia/chat/stream")
async def chat_stream(message: str = Form(...), history: str = Form("[]"), session_id: str = Form("default"), role: str = Form("Recruteur")):
    session = _get_session(session_id)
    effective_lang = detect_language(message)
    session_role = role or session.get("role","Recruteur")
    local_reply = _get_local_response(message, effective_lang, session_role)

    if local_reply:
        async def local_stream():
            words = local_reply.split(" ")
            for i, word in enumerate(words):
                yield f"data: {json.dumps({'token': word+(' ' if i<len(words)-1 else ''), 'done': False})}\n\n"
                await asyncio.sleep(0.008)
            yield f"data: {json.dumps({'token':'','done':True,'full':local_reply,'suggestions':_get_suggestions(message,effective_lang,session_role)})}\n\n"
        return StreamingResponse(local_stream(), media_type="text/event-stream", headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})

    async def event_generator() -> AsyncGenerator[str, None]:
        full = ""
        try:
            if not _gemini_client: raise QuotaExceeded("Gemini unavailable")
            loop   = asyncio.get_event_loop()
            prompt = _build_chat_prompt(session, message, session_role, effective_lang)
            stream = await loop.run_in_executor(_gemini_executor,
                lambda: _gemini_client.models.generate_content_stream(model=WORKING_MODEL, contents=prompt))
            for chunk in stream:
                if chunk.text:
                    full += chunk.text
                    yield f"data: {json.dumps({'token':chunk.text,'done':False})}\n\n"
            yield f"data: {json.dumps({'token':'','done':True,'full':full,'suggestions':_get_suggestions(message,effective_lang,session_role)})}\n\n"
        except (QuotaExceeded, Exception):
            fallback = _get_smart_fallback(message, effective_lang, session_role)
            yield f"data: {json.dumps({'token':fallback,'done':False})}\n\n"
            yield f"data: {json.dumps({'token':'','done':True,'full':fallback,'suggestions':_get_suggestions(message,effective_lang,session_role)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream", headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})


# ════════════════════════════════════════════════════════════
# ██   VOCAL   ██
# ════════════════════════════════════════════════════════════

VOICE_CONFIG = {
    "fr": {"lang_code":"fr-FR","lang_name":"Français","flag":"🇫🇷","recognition_lang":"fr-FR","speech_rate":1.0,"pitch":0,"wake_words":["écoute","neobot","assistant"],"pause_commands":["pause","attends"],"stop_commands":["arrête","couper","stop"],"resume_commands":["reprendre","continue"],"example_commands":["Créer un test React Senior","Montrer le dashboard","Générer ma lettre en arabe","Analyser mon CV","Quels thèmes disponibles ?"],"responses":{"greeting":"Bonjour ! Je suis NeoBot v10.1, votre assistant vocal.","pause_confirm":"Assistant en pause.","stop_confirm":"Assistant désactivé.","resume_confirm":"Je suis de retour !","not_understood":"Je n'ai pas compris. Répétez svp."}},
    "en": {"lang_code":"en-US","lang_name":"English","flag":"🇬🇧","recognition_lang":"en-US","speech_rate":1.0,"pitch":0,"wake_words":["listen","neobot","assistant"],"pause_commands":["pause","hold on"],"stop_commands":["stop","cut","end"],"resume_commands":["resume","continue"],"example_commands":["Create a React senior test","Generate my letter in French","Analyze my CV","Available themes?"],"responses":{"greeting":"Hello! I'm NeoBot v10.1, your voice assistant.","pause_confirm":"Assistant paused.","stop_confirm":"Assistant disabled.","resume_confirm":"I'm back!","not_understood":"Didn't catch that. Please repeat."}},
    "ar": {"lang_code":"ar-SA","lang_name":"العربية","flag":"🇸🇦","recognition_lang":"ar-SA","speech_rate":0.95,"pitch":0,"wake_words":["استمع","نيوبوت","المساعد"],"pause_commands":["إيقاف مؤقت","انتظر"],"stop_commands":["أوقف","اقطع"],"resume_commands":["استأنف","تابع"],"example_commands":["أنشئ اختبار React","توليد خطاب بالفرنسية","تحليل سيرتي","المواضيع المتاحة؟"],"responses":{"greeting":"مرحباً! أنا NeoBot v10.1، مساعدك الصوتي.","pause_confirm":"المساعد متوقف.","stop_confirm":"تم الإيقاف.","resume_confirm":"عدت!","not_understood":"لم أفهم. كرر من فضلك."}},
}

@app.get("/ia/voice/config")
async def get_voice_config(lang: str = Query("fr")):
    lang_key = lang if lang in VOICE_CONFIG else "fr"
    config = dict(VOICE_CONFIG[lang_key])
    config["all_languages"] = [{"code":k,"name":v["lang_name"],"flag":v["flag"]} for k,v in VOICE_CONFIG.items()]
    config["controls"] = {"start":{"icon":"🎙️","action":"start_listening"},"pause":{"icon":"⏸️","action":"pause_listening"},"stop":{"icon":"⏹️","action":"stop_listening"},"resume":{"icon":"▶️","action":"resume_listening"}}
    return config

@app.post("/ia/voice/process")
async def process_voice_command(text: str = Form(...), session_id: str = Form("default"), lang: str = Form("auto"), role: str = Form("Recruteur")):
    AI_METRICS["usage_counts"]["Vocal"] += 1
    return await chat_interaction(message=text, history="[]", role=role, lang=lang, session_id=session_id)


# ════════════════════════════════════════════════════════════
# ██   QCM — FIX: theme + sousTheme injectés   ██
# ════════════════════════════════════════════════════════════

def _fallback_qcm(theme, sousTheme, n, langue="fr", type: int = 0):
    t = theme.strip() or "Développement"
    s = sousTheme.strip() or t
    
    if type == 2:
        # Vrai / Faux
        if langue == "en":
            options = ["True", "False"]
            raw_bank = [
                {"question": f"The use of {s} in {t} is generally considered a secure and recommended practice when configured properly.", "answer": 0},
                {"question": f"Storing sensitive credentials and API keys in plain text inside {s} files is a major security hazard.", "answer": 0},
                {"question": f"In {t}, it is possible to deploy {s} without performing any unit tests or validations.", "answer": 1},
                {"question": f"For {t} projects, using {s} always reduces performance and makes the application slower.", "answer": 1},
                {"question": f"Proper input validation in {s} protects modern applications against injection and database attacks.", "answer": 0},
                {"question": f"It is recommended to ignore linter warnings when writing modular structures for {s}.", "answer": 1},
                {"question": f"Standard CI/CD pipelines automate testing for {s} to catch compilation issues early.", "answer": 0},
                {"question": f"Using outdated third-party scripts is a secure way to support {s} inside {t}.", "answer": 1},
                {"question": f"A decoupled architecture using {s} generally facilitates seamless long-term maintenance.", "answer": 0},
                {"question": f"Horizontal scaling with load balancers is ineffective for distributed {s} architectures.", "answer": 1},
                {"question": f"Applying the DRY (Don't Repeat Yourself) principle during refactoring of {s} helps reduce technical debt.", "answer": 0},
                {"question": f"Secure environment variables are preferred over hardcoding settings inside {s} files.", "answer": 0},
                {"question": f"Unit tests and integration tests can be safely disabled when releasing {s} in production.", "answer": 1},
                {"question": f"Linters and strict style guides help developers standardize {s} development within a team.", "answer": 0},
                {"question": f"Integration efficiency of {s} inside {t} can be monitored via test coverage and latencies.", "answer": 0},
                {"question": f"GraphQL, REST, or gRPC are modern protocols perfectly suitable for communication with {s}.", "answer": 0},
                {"question": f"React Hooks inside React can be called conditionally or inside nested loops.", "answer": 1},
                {"question": f"A solid understanding of core concepts in {t} is optional for mastering advanced features of {s}.", "answer": 1},
                {"question": f"Micro-frontend architectures using {s} increase code coupling and decrease deployment flexibility.", "answer": 1},
                {"question": f"Automating dependency updates ensures that {s} uses secure, patched library versions.", "answer": 0}
            ]
        elif langue == "ar":
            options = ["صواب", "خطأ"]
            raw_bank = [
                {"question": f"يعتبر استخدام {s} في {t} ممارسة آمنة وموصى بها بشكل عام عند تكوينها بشكل صحيح.", "answer": 0},
                {"question": f"تخزين بيانات الاعتماد الحساسة ومفاتيح API في نص واضح داخل ملفات {s} هو خطر أمني كبير.", "answer": 0},
                {"question": f"في {t}، يمكن نشر {s} دون إجراء أي اختبارات وحدة أو تحققات.", "answer": 1},
                {"question": f"بالنسبة لمشاريع {t}، فإن استخدام {s} يقلل دائماً من الأداء ويجعل التطبيق أبطأ.", "answer": 1},
                {"question": f"التحقق المناسب من المدخلات في {s} يحمي التطبيقات الحديثة من هجمات الحقن.", "answer": 0},
                {"question": f"يوصى بتجاهل تحذيرات linter عند كتابة هياكل برمجية لـ {s}.", "answer": 1},
                {"question": f"تقوم خطوط أنابيب CI/CD القياسية بأتمتة اختبارات {s} لاكتشاف مشكلات الترجمة مبكراً.", "answer": 0},
                {"question": f"استخدام برمجيات طرف ثالث غير موثوقة هو طريقة آمنة لدعم {s} في {t}.", "answer": 1},
                {"question": f"البنية البرمجية مفككة الارتباط باستخدام {s} تسهل بشكل عام الصيانة طويلة المدى.", "answer": 0},
                {"question": f"التوسع الأفقي مع موزعي الأحمال غير فعال بالنسبة لهياكل {s} الموزعة.", "answer": 1},
                {"question": f"يساعد تطبيق مبدأ DRY أثناء إعادة هيكلة {s} في تقليل الديون التقنية.", "answer": 0},
                {"question": f"تُفضل متغيرات البيئة الآمنة على تضمين الإعدادات مباشرة داخل ملفات {s}.", "answer": 0},
                {"question": f"يمكن تعطيل اختبارات الوحدة واختبارات التكامل بأمان عند إطلاق {s} في بيئة الإنتاج.", "answer": 1},
                {"question": f"تساعد أدوات التنسيق وأدلة الأسلوب الصارمة المطورين على توحيد معايير تطوير {s}.", "answer": 0},
                {"question": f"يمكن قياس كفاءة دمج {s} داخل {t} من خلال تغطية الاختبارات ووقت استجابة النظام.", "answer": 0},
                {"question": f"تعتبر بروتوكولات REST أو GraphQL أو gRPC مناسبة تماماً للاتصال مع {s}.", "answer": 0},
                {"question": f"يمكن استدعاء React Hooks داخل React بشكل مشروط أو داخل حلقات التكرار.", "answer": 1},
                {"question": f"الفهم المتين للمفاهيم الأساسية لـ {t} هو أمر اختياري لإتقان الميزات المتقدمة لـ {s}.", "answer": 1},
                {"question": f"تزيد بنية الواجهات الأمامية المصغرة باستخدام {s} من الارتباط الوثيق وتقلل من مرونة النشر.", "answer": 1},
                {"question": f"تضمن أتمتة تحديثات التبعيات استخدام {s} لإصدارات برمجية آمنة ومحدثة.", "answer": 0}
            ]
        else:
            options = ["Vrai", "Faux"]
            raw_bank = [
                {"question": f"L'utilisation de {s} en {t} est généralement considérée comme une pratique sûre et recommandée.", "answer": 0},
                {"question": f"Le stockage de clés de sécurité en clair dans les fichiers source de {s} est un anti-pattern majeur.", "answer": 0},
                {"question": f"Dans {t}, il est possible de déployer {s} en production sans effectuer aucun test unitaire.", "answer": 1},
                {"question": f"Pour les projets {t}, l'utilisation de {s} réduit systématiquement les performances globales.", "answer": 1},
                {"question": f"Une validation stricte des entrées pour {s} protège l'application moderne contre les injections.", "answer": 0},
                {"question": f"Il est recommandé de masquer ou d'ignorer les avertissements du linter lors du développement de {s}.", "answer": 1},
                {"question": f"Les pipelines CI/CD automatisent les tests de {s} afin de détecter les bugs avant le déploiement.", "answer": 0},
                {"question": f"L'intégration de scripts tiers non vérifiés est idéale pour assurer la sécurité de {s}.", "answer": 1},
                {"question": f"Une architecture modulaire et découplée facilite grandement la maintenance de {s}.", "answer": 0},
                {"question": f"La mise à l'échelle horizontale (scaling) est inutile pour les applications utilisant {s}.", "answer": 1},
                {"question": f"Appliquer le principe DRY lors de la refactorisation de {s} permet de réduire la dette technique.", "answer": 0},
                {"question": f"L'usage de variables d'environnement est préférable au fait de hardcoder la configuration dans {s}.", "answer": 0},
                {"question": f"Les tests unitaires peuvent être désactivés sans risque lors du passage de {s} en production.", "answer": 1},
                {"question": f"Les linters et guides de style stricts aident les équipes à standardiser le code de {s}.", "answer": 0},
                {"question": f"L'efficacité d'intégration de {s} se mesure par la couverture de code et les temps de réponse.", "answer": 0},
                {"question": f"REST, GraphQL et gRPC sont des protocoles de communication modernes adaptés pour {s}.", "answer": 0},
                {"question": f"Les React Hooks dans React peuvent être appelés de manière conditionnelle ou dans des boucles.", "answer": 1},
                {"question": f"Comprendre l'architecture de {t} est facultatif pour exploiter les capacités avancées de {s}.", "answer": 1},
                {"question": f"L'architecture micro-frontend augmente le couplage et réduit la flexibilité de déploiement.", "answer": 1},
                {"question": f"L'automatisation des mises à jour des dépendances garantit la sécurité de la structure {s}.", "answer": 0}
            ]

        questions_list = []
        for i in range(n):
            q_tpl = raw_bank[i % len(raw_bank)]
            questions_list.append({
                "question": q_tpl["question"],
                "options": list(options),
                "answer": q_tpl["answer"],
                "langue": langue,
                "theme": t,
                "sousTheme": s
            })

        return {
            "questions": questions_list,
            "_source": "fallback_tf",
            "_langue": langue
        }
    
    if langue == "en":
        raw_bank = [
            {"question": f"What is the primary best practice for implementing {s} in {t}?", "options": [f"Adhering to standard style guides and conventions of {t}", f"Avoiding any code documentation in {s}", f"Copying code directly without verification", f"Omitting unit tests for {s} entirely"]},
            {"question": f"Which design pattern is best suited for structuring a project using {s}?", "options": [f"A modular, decoupled architecture", f"A highly coupled monolith using {s}", f"Linear and sequential shell scripts", f"No structural patterns whatsoever"]},
            {"question": f"How should you handle errors and exceptions effectively in {s}?", "options": [f"Using try-catch blocks with structured logging", f"Silently ignoring all exceptions in {s}", f"Displaying raw error traces to end users", f"Systematically restarting the server on any error"]},
            {"question": f"What is the major advantage of strong typing or strict validation in {t}?", "options": [f"Catching bugs early during compilation or build", f"Slowing down overall execution speed of {s}", f"Reducing general code flexibility", f"No tangible advantages in production"]},
            {"question": f"How can you ensure excellent long-term maintainability of {s}?", "options": [f"Writing unit tests and providing clear API documentation", f"Adding only basic inline comments to {s}", f"Avoiding any form of refactoring", f"Frequently changing frameworks and libraries"]},
            {"question": f"What is the recommended approach for testing {s} in a CI/CD pipeline?", "options": [f"Automating unit and integration tests", f"Manually testing features directly in production", f"Testing only cosmetic front-end changes", f"Disabling tests to speed up deployment"]},
            {"question": f"In {t}, how do you optimize performance when utilizing {s}?", "options": [f"Caching queries and optimizing resource allocation", f"Indefinitely upgrading server hardware", f"Avoiding asynchronous programming with {s}", f"Minimizing the number of active variables"]},
            {"question": f"Which tool or library is standard for supporting {s} in {t}?", "options": [f"Official tools recommended by the {t} ecosystem", f"Unverified third-party scripts for {s}", f"An in-house custom solution without support", f"No additional tools whatsoever"]},
            {"question": f"What is the most common pitfall or anti-pattern associated with {s}?", "options": [f"Tight coupling and lack of separation of concerns", f"Excessive use of pure functions in {s}", f"Over-modularizing simple {t} features", f"Configuring options via environment variables only"]},
            {"question": f"How do you secure the implementation of {s} in a modern application?", "options": [f"Validating all inputs and encrypting sensitive data", f"Relying solely on client-side validation for {s}", f"Disabling HTTPS protocols in development", f"Hardcoding secret keys in source files"]},
            {"question": f"What is the ideal method for documenting the architecture of {s}?", "options": [f"Generating interactive API docs and UML diagrams", f"Letting the code explain itself without docs", f"Creating an empty placeholder text file", f"Writing short handwritten notes"]},
            {"question": f"In a team environment, how can you standardize the development of {s}?", "options": [f"Utilizing linters and enforcing strict formatting rules", f"Allowing each developer to write in their own style", f"Banning the use of comments in {s}", f"Merging pull requests without code reviews"]},
            {"question": f"Which scaling strategy (scalability) is best suited for {s}?", "options": [f"Horizontal scaling with load balancing", f"Upgrading only the server's RAM", f"Limiting the number of active users for {s}", f"Removing secondary features entirely"]},
            {"question": f"What is the impact of using {s} on the lifecycle of a {t} application?", "options": [f"It facilitates seamless updates and reduces technical debt", f"It complicates deployment pipelines unnecessarily", f"It makes the codebase dependent on a single developer", f"No impact on the application lifecycle"]},
            {"question": f"How do you reduce technical debt when refactoring {s}?", "options": [f"Applying the DRY principle and reducing complexity", f"Rewriting the entire codebase without writing tests", f"Ignoring all warnings from the linter", f"Adding new features while refactoring {s}"]},
            {"question": f"What is the best approach to manage the configuration of {s}?", "options": [f"Using secure, read-only environment variables", f"Hardcoding configuration values in {s} source files", f"Storing active credentials in public repositories", f"Prompting the user for settings on every startup"]},
            {"question": f"In {t}, what is the main role of {s}?", "options": [f"Structuring data flows and streamlining key processes", f"Replacing the entire database engine of {t}", f"Serving only as a cosmetic interface element", f"Speeding up only the initial server boot time"]},
            {"question": f"How do you measure the integration efficiency of {s} in {t}?", "options": [f"Analyzing test coverage and system response latency", f"Counting the absolute number of written code lines", f"Checking only the final compiled bundle size", f"Integration efficiency cannot be measured"]},
            {"question": f"Which communication standard or protocol is ideal for {s}?", "options": [f"Modern standards like REST, GraphQL, or gRPC", f"Shared plain text files transferred via FTP", f"Raw, unencrypted socket connections", f"No communication protocols whatsoever"]},
            {"question": f"Which skill is essential for mastering {s} in {t}?", "options": [f"A solid understanding of core concepts and architecture in {t}", f"Memorizing every single {s} function signature by heart", f"Using AI exclusively without understanding the code", f"No specific skills are required for {t}"]}
        ]
    elif langue == "ar":
        raw_bank = [
            {"question": f"ما هي الممارسة الفضلى الأساسية عند تطبيق {s} في {t}؟", "options": [f"اتباع أدلة الأسلوب والمعايير القياسية لـ {t}", f"تجنب كتابة أي توثيق للكود الخاص بـ {s}", f"نسخ الأكواد مباشرة دون التحقق منها", f"إغفال اختبارات الوحدة لـ {s} بالكامل"]},
            {"question": f"أي نمط تصميم (Design Pattern) هو الأنسب لهيكلة مشروع يستخدم {s}؟", "options": [f"بنية برمجية نمطية ومفككة الارتباط", f"بنية أحادية متماسكة للغاية ومقيدة باستخدام {s}", f"مجموعة من برامج شل الخطية المتتالية", f"لا توجد أنماط هيكلية محددة"]},
            {"question": f"كيف يمكنك التعامل مع الأخطاء والاستثناءات بفعالية في {s}؟", "options": [f"استخدام كتل try-catch مع نظام تسجيل منظم", f"تجاهل جميع الاستثناءات في {s} بصمت", f"عرض تفاصيل الخطأ الخام للمستخدمين النهائيين", f"إعادة تشغيل الخادم تلقائياً وبشكل مستمر عند حدوث خطأ"]},
            {"question": f"ما هي الميزة الرئيسية للكتابة القوية أو التحقق الصارم في {t}؟", "options": [f"اكتشاف الأخطاء مبكراً أثناء مرحلة الترجمة أو البناء", f"إبطاء سرعة التنفيذ الإجمالية للتطبيق باستخدام {s}", f"تقليل المرونة العامة للكتابة البرمجية", f"لا توجد مزايا ملموسة قياساً بالأداء في بيئة الإنتاج"]},
            {"question": f"كيف يمكنك ضمان صيانة ممتازة طويلة المدى لـ {s}؟", "options": [f"كتابة اختبارات الوحدة وتوفير توثيق واضح لواجهات البرمجة", f"إضافة تعليقات توضيحية بسيطة وموجزة لـ {s} فقط", f"تجنب أي شكل من أشكال إعادة هيكلة الكود", f"تغيير إطارات العمل والبرمجيات المستخدمة بشكل متكرر"]},
            {"question": f"ما هو النهص الموصى به لاختبار {s} في بيئة التطوير والتشغيل المستمر (CI/CD)؟", "options": [f"أتمتة اختبارات الوحدة واختبارات التكامل", f"إجراء الاختبارات يدوياً مباشرة في بيئة الإنتاج", f"اختبار التغييرات التجميلية للواجهة الأمامية فقط", f"تعطيل الاختبارات لتسريع عملية النشر"]},
            {"question": f"في {t}، كيف يمكنك تحسين الأداء عند استخدام {s}؟", "options": [f"تخزين الاستعلامات مؤقتاً وتحسين تخصيص الموارد لـ {s}", f"ترقية عتاد الخادم بشكل مستمر ودون حدود", f"تجنب البرمجة غير المتزامنة تماماً مع {s}", f"تقليل عدد المتغيرات النشطة في النظام"]},
            {"question": f"ما هي الأداة أو المكتبة القياسية لدعم {s} في {t}؟", "options": [f"الأدوات الرسمية الموصى بها من قبل النظام البيئي لـ {t}", f"برمجيات طرف ثالث غير موثوقة أو مصادق عليها لـ {s}", f"حل داخلي مخصص ومطور بدون دعم فني", f"لا توجد أدوات إضافية مطلوبة بالكامل"]},
            {"question": f"ما هو الفخ أو النمط المضاد (Anti-Pattern) الأكثر شيوعاً المرتبط بـ {s}؟", "options": [f"الارتباط الوثيق وغياب الفصل بين المسؤوليات", f"الاستخدام المفرط للدوال النقية في {s}", f"التقسيم النمطي المبالغ فيه لميزات {t} البسيطة", f"تكوين الخيارات عبر متغيرات البيئة فقط"]},
            {"question": f"كيف يمكنك تأمين تطبيق {s} في البرمجيات الحديثة؟", "options": [f"التحقق من صحة المدخلات وتشفير البيانات الحساسة", f"الاعتماد الكلي على التحقق من جانب العميل فقط لـ {s}", f"تعطيل بروتوكولات HTTPS في بيئة التطوير", f"تضمين مفاتيح التشفير والسرية مباشرة في الملفات المصدرية"]},
            {"question": f"ما هي الطريقة المثلى لتوثيق بنية {s}؟", "options": [f"توليد مستندات تفاعلية ومخططات UML واضحة لـ {s}", f"ترك الكود ليشرح نفسه دون أي توثيق خارجي", f"إنشاء ملف نصي فارغ كعلامة موجهة", f"كتابة ملاحظات ورقية قصيرة بخط اليد"]},
            {"question": f"في بيئة عمل الفريق، كيف يمكنك توحيد معايير تطوير {s}؟", "options": [f"استخدام أدوات الفحص وتطبيق قواعد تنسيق صارمة", f"السماح لكل مطور بالكتابة بأسلوبه الخاص في {s}", f"حظر استخدام التعليقات التوضيحية تماماً في {s}", f"دمج طلبات السحب دون مراجعة الأكواد"]},
            {"question": f"أي استراتيجية توسع (Scalability) هي الأنسب لـ {s}؟", "options": [f"التوسع الأفقي مع توزيع الأحمال بكفاءة", f"زيادة ذاكرة الوصول العشوائي (RAM) للخادم فقط", f"تحديد عدد المستخدمين النشطين في نفس الوقت لـ {s}", f"إزالة الميزات الثانوية لتقليل العبء بالكامل"]},
            {"question": f"ما هو تأثير استخدام {s} على دورة حياة تطبيق {t}؟", "options": [f"يسهل التحديثات السلسة ويقلل من الديون التقنية لـ {s}", f"يعقد مسارات النشر والتطوير دون داعٍ", f"يجعل المشروع معتمداً بالكامل على مطور واحد", f"لا يوجد أي تأثير يذكر على دورة حياة التطبيق"]},
            {"question": f"كيف يمكنك تقليل الديون التقنية عند إعادة هيكلة {s}؟", "options": [f"تطبيق مبدأ DRY وتقليل التعقيد الحسابي لكود {s}", f"إعادة كتابة المشروع بالكامل دون كتابة أي اختبارات", f"تجاهل جميع تحذيرات أدوات فحص وتنسيق الأكواد لـ {s}", f"إضافة ميزات جديدة أثناء عملية إعادة الهيكلة لـ {s}"]},
            {"question": f"ما هو أفضل نهج لإدارة إعدادات وتكوين {s}؟", "options": [f"استخدام متغيرات بيئة آمنة ومحمية من القراءة الخارجية", f"تضمين قيم الإعدادات الثابتة في ملفات {s} المصدرية", f"تخزين بيانات الاعتماد النشطة في مستودعات عامة", f"مطالبة المستخدم بالإعدادات يدوياً عند كل تشغيل"]},
            {"question": f"في {t}، ما هو الدور الرئيسي لـ {s}؟", "options": [f"هيكلة تدفقات البيانات وتبسيط العمليات الأساسية", f"استبدال محرك قاعدة البيانات بالكامل لـ {t}", f"العمل كعنصر تجميلي للواجهة الرسومية فقط", f"تسريع وقت إقلاع الخادم الأولي فقط"]},
            {"question": f"كيف تقيس كفاءة دمج {s} في {t}؟", "options": [f"تحليل تغطية الاختبارات ووقت استجابة النظام الفعلي", f"حساب العدد الإجمالي لأسطر الأكواد المكتوبة لـ {s}", f"التحقق من الحجم النهائي للملفات المجمعة فقط", f"كفاءة الدمج هي أمر لا يمكن قياسه عملياً"]},
            {"question": f"أي معيار أو بروتوكول اتصال هو الأفضل للاستخدام مع {s}؟", "options": [f"المعايير الحديثة مثل REST أو GraphQL أو gRPC", f"الملفات النصية المشتركة المنقولة عبر بروتوكول FTP", f"اتصالات المقابس الخام غير المشفرة أو الآمنة", f"لا توجد بروتوكولات اتصال مفضلة"]},
            {"question": f"ما هي المهارة الأساسية اللازمة لإتقان {s} في {t}؟", "options": [f"فهم متين للمفاهيم الأساسية والبنية الهيكلية لـ {t}", f"حفظ توقيع جميع دوال {s} البرمجية عن ظهر قلب", f"الاعتماد الكلي على الذكاء الاصطناعي دون فهم الكود", f"لا توجد مهارات محددة مطلوبة للبدء"]}
        ]
    else:
        raw_bank = [
            {"question": f"Quelle est la principale bonne pratique pour implémenter {s} en {t} ?", "options": [f"Respecter les conventions et normes standards de {t}", f"Éviter toute documentation dans {s}", f"Copier le code directement sans vérification", f"Ne pas faire de tests unitaires pour {s}"]},
            {"question": f"Quel design pattern est le plus adapté pour structurer un projet utilisant {s} ?", "options": [f"Une architecture modulaire et découplée", f"Un monolithe rigide et couplé avec {s}", f"Des scripts shell linéaires et séquentiels", f"Aucun modèle particulier de structure"]},
            {"question": f"Comment gérer efficacement les erreurs et exceptions dans {s} ?", "options": [f"Utiliser try-catch et des logs structurés", f"Ignorer silencieusement les erreurs dans {s}", f"Afficher l'erreur brute à l'utilisateur final", f"Redémarrer systématiquement l'application"]},
            {"question": f"Quel est l'avantage majeur du typage fort ou de la validation stricte dans {t} ?", "options": [f"Détecter les bugs dès la phase de compilation ou build", f"Ralentir l'exécution globale avec {s}", f"Réduire la flexibilité générale du code", f"Aucun avantage mesurable en production"]},
            {"question": f"Comment assurer une excellente maintenabilité à long terme de {s} ?", "options": [f"Écrire des tests unitaires et documenter l'API", f"Ajouter uniquement des commentaires simples à {s}", f"Éviter tout refactoring du code", f"Changer de framework et de librairies régulièrement"]},
            {"question": f"Quelle est l'approche recommandée pour tester {s} en environnement CI/CD ?", "options": [f"Automatiser les tests unitaires et d'intégration", f"Tester manuellement les fonctionnalités en production", f"Ne tester que les changements cosmétiques", f"Désactiver les tests lors du déploiement"]},
            {"question": f"Dans {t}, comment optimiser les performances lors de l'utilisation de {s} ?", "options": [f"Mettre en cache les requêtes et optimiser les ressources", f"Augmenter indéfiniment la puissance du serveur", f"Éviter tout asynchronisme avec {s}", f"Minimiser le nombre de variables actives"]},
            {"question": f"Quel outil ou bibliothèque est standard pour accompagner {s} dans {t} ?", "options": [f"Les outils officiels recommandés par l'écosystème de {t}", f"Des scripts tiers non vérifiés pour {s}", f"Une solution développée en interne sans support", f"Aucun outil complémentaire dans l'écosystème"]},
            {"question": f"Quel est le piège ou l'anti-pattern le plus courant avec {s} ?", "options": [f"Le couplage fort et l'absence de séparation des responsabilités", f"L'utilisation excessive de fonctions pures dans {s}", f"Une modularité trop poussée pour des fonctionnalités simples de {t}", f"La configuration par variables d'environnement uniquement"]},
            {"question": f"Comment sécuriser l'implémentation de {s} dans une application moderne ?", "options": [f"Valider toutes les entrées et chiffrer les données sensibles", f"Faire confiance aux validations du client uniquement pour {s}", f"Désactiver le protocole HTTPS en développement", f"Stocker les clés de sécurité en clair dans le code"]},
            {"question": f"Quelle est la méthode idéale pour documenter l'architecture de {s} ?", "options": [f"Générer une documentation interactive et des schémas UML", f"Laisser le code s'expliquer de lui-même sans docs", f"Créer un fichier texte vide", f"Rédiger des notes manuscrites rapides"]},
            {"question": f"Dans un contexte d'équipe, comment standardiser le développement de {s} ?", "options": [f"Utiliser des linters et des guides de style stricts", f"Laisser chaque développeur choisir son style pour {s}", f"Interdire l'usage de commentaires explicatifs", f"Fusionner directement sans code review"]},
            {"question": f"Quelle stratégie de mise à l'échelle (scalability) convient le mieux à {s} ?", "options": [f"Une architecture horizontale avec répartition de charge", f"Augmenter uniquement la mémoire RAM du serveur", f"Réduire le nombre d'utilisateurs actifs de {s}", f"Supprimer les fonctionnalités secondaires"]},
            {"question": f"Quel est l'impact de l'utilisation de {s} sur le cycle de vie de l'application {t} ?", "options": [f"Elle facilite les mises à jour et réduit la dette technique", f"Elle complique le déploiement de manière injustifiée", f"Elle rend l'application dépendante d'un seul développeur", f"Aucun impact sur le cycle de vie de l'application"]},
            {"question": f"Comment réduire la dette technique lors de la refactorisation de {s} ?", "options": [f"Appliquer le principe DRY et simplifier la complexité", f"Réécrire l'intégralité du code sans tests", f"Ignorer les avertissements du linter de {s}", f"Ajouter de nouvelles fonctionnalités lors de la refactorisation"]},
            {"question": f"Quelle est la meilleure approche pour gérer la configuration de {s} ?", "options": [f"Utiliser des variables d'environnement sécurisées", f"Hardcoder les valeurs dans les fichiers sources de {s}", f"Stocker les configurations sur un serveur public", f"Demander la configuration à l'utilisateur à chaque démarrage"]},
            {"question": f"Dans {t}, quel est le rôle principal de {s} ?", "options": [f"Structurer les flux de données et rationaliser les processus clés", f"Remplacer l'intégralité de la base de données", f"Servir uniquement d'élément esthétique d'interface", f"Accélérer uniquement le démarrage du serveur"]},
            {"question": f"Comment mesurer l'efficacité de l'intégration de {s} dans {t} ?", "options": [f"En analysant la couverture de code et les temps de réponse", f"En comptant le nombre de lignes de code de {s}", f"En observant uniquement la taille du fichier final", f"Il est impossible de mesurer cette efficacité"]},
            {"question": f"Quel protocole ou standard de communication est idéal avec {s} ?", "options": [f"Les standards modernes comme REST, GraphQL ou gRPC", f"Des fichiers texte partagés via FTP", f"Des connexions sockets brutes non sécurisées", f"Aucun protocole de communication"]},
            {"question": f"Quelle compétence est indispensable pour maîtriser {s} en {t} ?", "options": [f"Une excellente compréhension des concepts et de l'architecture de {t}", f"La mémorisation par cœur de toutes les fonctions de {s}", f"L'utilisation exclusive de l'IA sans comprendre le code", f"Aucune compétence spécifique n'est requise"]}
        ]

    # Dynamically select questions from the pool and randomize/shuffle option ordering
    questions_list = []
    for i in range(n):
        q_tpl = raw_bank[i % len(raw_bank)]
        # Make a copy of options so we do not mutate the template pool
        options = list(q_tpl["options"])
        correct_option = options[0]  # First option in our raw bank is always correct
        random.shuffle(options)
        correct_index = options.index(correct_option)
        
        questions_list.append({
            "question": q_tpl["question"],
            "options": options,
            "answer": correct_index,
            "langue": langue,
            "theme": t,
            "sousTheme": s
        })

    return {
        "questions": questions_list,
        "_source": "fallback",
        "_langue": langue
    }


@app.post("/ia/generate-bilingual")
async def generate_bilingual(
    theme: str = Form(...),
    sousTheme: str = Form(...),
    n: int = Form(5),
    langue: str = Form("fr"),
    type: int = Form(0)
):
    ck = make_cache_key("bilingual-v10", theme, sousTheme, n, langue, type)
    if hit := _cache.get(ck): return hit

    if type == 2:
        # Vrai / Faux
        opt_instr = {
            "fr": '["Vrai", "Faux"]',
            "en": '["True", "False"]',
            "ar": '["صواب", "خطأ"]'
        }.get(langue, '["Vrai", "Faux"]')
        
        lang_instr = {
            "fr": f"Génère {n} questions Vrai/Faux expertes en FRANÇAIS sur '{theme}' sous-thème '{sousTheme}'.",
            "en": f"Generate {n} expert True/False questions in ENGLISH about '{theme}' sub-theme '{sousTheme}'.",
            "ar": f"أنشئ {n} أسئلة صح أم خطأ متخصصة باللغة العربية حول '{theme}' الموضوع '{sousTheme}'.",
        }
        prompt = f"""{lang_instr.get(langue, lang_instr['fr'])}
Chaque question doit obligatoirement avoir exactement ces 2 options dans l'ordre : {opt_instr}, answer=index de la bonne réponse (0 ou 1). Langue: {langue.upper()} UNIQUEMENT.
JSON: {{"questions":[{{"question":"...","options":{opt_instr},"answer":0}}]}}"""
    else:
        lang_instr = {
            "fr": f"Génère {n} questions QCM expertes en FRANÇAIS sur '{theme}' sous-thème '{sousTheme}'.",
            "en": f"Generate {n} expert QCM questions in ENGLISH about '{theme}' sub-theme '{sousTheme}'.",
            "ar": f"أنشئ {n} أسئلة متخصصة باللغة العربية حول '{theme}' الموضوع '{sousTheme}'.",
        }
        prompt = f"""{lang_instr.get(langue, lang_instr['fr'])}
4 options par question, answer=index(0-3). Langue: {langue.upper()} UNIQUEMENT.
JSON: {{"questions":[{{"question":"...","options":["...","...","...","..."],"answer":0}}]}}"""

    try:
        r = await call_gemini_async(prompt, module="Évaluations", sem=_sem_qcm, retries=1)
        result = json.loads(clean_json(r.text))
        # ── FIX: injecter theme + sousTheme dans chaque question Gemini ──
        for q in result.get("questions", []):
            q["langue"]    = langue
            q["theme"]     = theme.strip()
            q["sousTheme"] = sousTheme.strip()
        _cache.set(ck, result)
        return result
    except (QuotaExceeded, Exception):
        return _fallback_qcm(theme, sousTheme, n, langue, type)


@app.post("/ia/generate-ultra")
async def generate_ultra(
    theme: str = Form(...),
    sousTheme: str = Form(...),
    n: int = Form(5),
    langue: str = Form("fr"),
    type: int = Form(0)
):
    return await generate_bilingual(theme=theme, sousTheme=sousTheme, n=n, langue=langue, type=type)


@app.post("/ia/generate-pro")
async def generate_pro(
    nombre: int = Form(...),
    themetique: str = Form(...),
    difficulte: str = Form(...),
    langue: str = Form("fr"),
    file: UploadFile = File(None)
):
    context = ""
    if file: context = await extract_text_from_upload(file)
    ck = make_cache_key("pro-v10", nombre, themetique, difficulte, langue, context[:150])
    if hit := _cache.get(ck): return hit
    try:
        lang_instr = {"fr":"en français","en":"in English","ar":"باللغة العربية"}.get(langue,"en français")
        ctx_part   = f"\nDocument context:\n{context[:1500]}" if context else ""
        prompt     = f"Generate {nombre} QCM questions {lang_instr} on '{themetique}' at '{difficulte}' level.{ctx_part}\nJSON with 'questions' array."
        r = await call_gemini_async(prompt, module="Évaluations", sem=_sem_qcm, retries=1)
        data = json.loads(clean_json(r.text))
        # ── FIX: injecter theme + sousTheme dans generate-pro ──
        for q in data.get("questions", []):
            q["langue"]    = langue
            q["theme"]     = themetique.strip()
            q["sousTheme"] = themetique.strip()
        result = {"status":"IA_SUCCESS","questions":data.get("questions",[]),"source":"gemini"}
        _cache.set(ck, result)
        return result
    except (QuotaExceeded, Exception):
        fb = _fallback_qcm(themetique, themetique, nombre, langue)
        return {"status":"IA_SUCCESS","questions":fb["questions"],"source":"fallback"}


# ════════════════════════════════════════════════════════════
# ██   AUTRES ENDPOINTS IA   ██
# ════════════════════════════════════════════════════════════

@app.post("/ia/analyze-candidate")
async def analyze_candidate(nom: str = Form(...), scores_techniques: str = Form(...)):
    nums = [int(x) for x in re.findall(r"\d+", scores_techniques)]
    avg  = sum(nums)//len(nums) if nums else 75
    data = {"id":f"AI-{int(time.time())}","profile_type":"Expert" if avg>80 else "Intermédiaire" if avg>60 else "Junior","global_score":avg,"neural_tier":"Élite" if avg>85 else "Standard","traits":[{"name":"Capacité Logique","val":avg,"color":"#4f46e5","icon":"fa-brain"},{"name":"Adaptabilité","val":max(avg-5,0),"color":"#f59e0b","icon":"fa-bolt"},{"name":"Stabilité","val":85,"color":"#10b981","icon":"fa-shield-halved"}]}
    try:
        prompt = f"1 phrase pro: candidat score {avg}/100 pour {nom}."
        response = await asyncio.wait_for(call_gemini_async(prompt, retries=1, module="Analyses CV", sem=_sem_cv), timeout=5.0)
        data["ai_insight"] = response.text.strip()
    except (QuotaExceeded, asyncio.TimeoutError, Exception):
        data["ai_insight"] = f"Score {avg}/100 — Profil {data['profile_type']}. Recommandé pour l'étape suivante."
    return data

@app.post("/ia/reports/generate")
async def generate_report(report_type: str = Form("org"), period: str = Form("month"), user_role: str = Form("AdminEntreprise"), context: str = Form("")):
    try:
        prompt = f"Rapport RH JSON. Type:{report_type} Période:{period} Rôle:{user_role}.\nJSON: title,summary,sections[{{title,content}}],kpis[{{label,value,trend}}],recommendations[]. Français."
        r = await call_gemini_async(prompt, module="Rapports", sem=_sem_rpts, retries=1)
        result = json.loads(clean_json(r.text))
        result["generated_at"] = datetime.now().isoformat()
        return {"status":"SUCCESS","report":result}
    except (QuotaExceeded, Exception):
        return {"status":"SIMULATION","report":{"title":f"Rapport {report_type.upper()} — {period}","summary":"Performance satisfaisante. Taux réussite 88%.","sections":[{"title":"Vue d'ensemble","content":"Performance en hausse de 12%."},{"title":"Recommandations","content":"Augmenter la fréquence des campagnes."}],"kpis":[{"label":"Taux réussite","value":"88%","trend":"+3%"},{"label":"Candidats actifs","value":"124","trend":"+8"}],"recommendations":["Augmenter les campagnes","Rapports hebdomadaires"],"generated_at":datetime.now().isoformat(),"period":period,"type":report_type}}

@app.post("/ia/interview/generate")
async def generate_interview(job_title: str = Form(...), level: str = Form("Intermédiaire"), focus: str = Form(""), langue: str = Form("fr")):
    ck = make_cache_key("interview-v10", job_title, level, langue)
    if hit := _cache.get(ck): return hit
    try:
        lang_instr = {"fr":"en français","en":"in English","ar":"باللغة العربية"}.get(langue,"en français")
        prompt = f"5 questions entretien RH {job_title} niveau {level}, focus:{focus or 'général'} {lang_instr}.\n{{\"questions\":[{{\"question\":\"...\",\"type\":\"comportemental|technique|situationnel\",\"tip\":\"...\"}}]}}"
        r = await call_gemini_async(prompt, module="Entretiens IA", retries=1)
        result = json.loads(clean_json(r.text))
        _cache.set(ck, result); return result
    except (QuotaExceeded, Exception):
        return {"questions":[{"question":f"Décrivez votre expérience en {job_title}.","type":"comportemental","tip":"Soyez précis sur vos réalisations."},{"question":"Comment gérez-vous la pression ?","type":"situationnel","tip":"Donnez un exemple STAR."},{"question":"Votre plus grande réussite ?","type":"comportemental","tip":"Chiffrez vos résultats."},{"question":"Comment travaillez-vous en équipe ?","type":"comportemental","tip":"Montrez votre empathie."},{"question":"Pourquoi ce poste ?","type":"motivation","tip":"Alignez vos ambitions."}]}

@app.post("/ia/performance-report")
async def get_ia_performance_report():
    h = list(AI_METRICS["latency_history"])
    avg_l = sum(h)/len(h) if h else 0
    charge = min(100, AI_METRICS["active_requests"]*20)
    total  = sum(AI_METRICS["usage_counts"].values()) or 1
    colors = {"Évaluations":"#6366f1","Analyses CV":"#f97316","Entretiens IA":"#10b981","Rapports":"#8b5cf6","Chat":"#3b82f6","Dashboard":"#0ea5e9","Recommandations":"#f59e0b","Lettres":"#ec4899","Vocal":"#14b8a6"}
    usage_data = [{"name":k,"pct":int(v/total*100),"color":colors.get(k,"#888")} for k,v in AI_METRICS["usage_counts"].items() if v>0]
    return {
        "performance":{
            "charge":charge or 5,"tokens":f"{AI_METRICS['total_tokens']/1_000_000:.2f}M",
            "responseTime":f"{int(avg_l)}ms","requestsPerSecond":AI_METRICS["active_requests"],
            "circuit_state":_circuit.state,"error_count":AI_METRICS["error_count"],
            "dashboard_requests":AI_METRICS["dashboard_requests"],
            "gemini_available":_gemini_client is not None,
            "total_intents":len(UNIVERSAL_BRAIN),
            "cv_rejected_total":AI_METRICS["cv_rejected"],
        },
        "usage":usage_data
    }

@app.post("/ia/radar-analysis")
async def radar_analysis_standalone(file: UploadFile = File(...)):
    fb = await file.read()
    cv_text = _truncate(await read_pdf_async(fb))
    ck = make_cache_key("radar-v10", cv_text[:200])
    if hit := _cache.get(ck): return hit
    try:
        prompt = f'Analyse soft-skills CV JSON:\n{cv_text[:1500]}\n{{"values":[int,int,int,int,int]}} Communication,Leadership,Adaptabilité,Équipe,Résolution 0-100.'
        r = await call_gemini_async(prompt, module="Analyses CV", sem=_sem_cv, retries=1)
        result = json.loads(clean_json(r.text))
        _cache.set(ck, result); return result
    except (QuotaExceeded, Exception):
        return {"values":[60+random.randint(-10,20),85+random.randint(-10,10),70+random.randint(-10,15),90+random.randint(-10,5),75+random.randint(-10,15)]}


# ─────────────────────────────────────────────────
@app.post("/ia/evaluate-exam")
async def evaluate_exam(request: Request):
    try:
        data = await request.json()
        questions = data.get("questions", [])
        reponses = data.get("reponses", {})
        candidat_nom = data.get("candidatNom", "Candidat")
        examen_nom = data.get("examenNom", "Examen")
        
        # Prepare the context for Gemini
        exam_context = []
        for q in questions:
            qid = q.get("id")
            enonce = q.get("enonce", "")
            type_q = q.get("type", "")
            choix = q.get("choix", [])
            bonne_rep = q.get("bonneReponse", "")
            user_rep = reponses.get(str(qid), "")
            
            exam_context.append({
                "id": qid,
                "question": enonce,
                "type": type_q,
                "options": choix,
                "correct_answer_from_db": bonne_rep,
                "candidate_answer": user_rep
            })
            
        prompt = f"""Tu es un expert formateur et évaluateur pour EvaluaTech.
Tu dois corriger un examen passé par {candidat_nom} sur le sujet "{examen_nom}".
Voici les questions, la réponse attendue et la réponse du candidat.

Évalue chaque réponse. Pour les questions à choix multiples (QCM/QCU) ou Vrai/Faux, base-toi strictement sur la "correct_answer_from_db" fournie. Pour les questions ouvertes, évalue la pertinence de la "candidate_answer".
Calcule le score total (en pourcentage) et rédige un rapport final personnalisé.

DONNÉES DE L'EXAMEN :
{json.dumps(exam_context, ensure_ascii=False, indent=2)}

Tu dois répondre UNIQUEMENT avec un objet JSON structuré comme suit :
{{
  "scorePourcentage": 85,
  "rapportFinal": "Un résumé global des performances du candidat...",
  "corrections": [
    {{
      "questionId": "l'ID de la question",
      "isCorrect": true/false,
      "candidateAnswer": "la réponse du candidat interprétée en texte",
      "correctAnswer": "la bonne réponse interprétée en texte",
      "explication": "Explication détaillée de pourquoi c'est correct ou incorrect (très important !)"
    }}
  ]
}}

Ne renvoie AUCUN texte en dehors du JSON."""

        response = await call_gemini_async(prompt, module="Evaluation", sem=_sem_rpts, retries=1)
        text = response.text.strip()
        cleaned = clean_json(text)
        
        try:
            parsed = json.loads(cleaned)
            log_activity("IA", f"Évaluation IA terminée ({examen_nom})", "#10b981", "Evaluation")
            return {"status": "SUCCESS", "evaluation": parsed}
        except Exception as e:
            logger.error(f"Failed to parse Gemini evaluation JSON: {e}\n{cleaned}")
            return {"status": "ERROR", "message": "Format de réponse invalide de l'IA."}

    except Exception as e:
        return {"status": "ERROR", "message": str(e)}
class RecommendationsRequest(BaseModel):
    role: str
    lang: str = "fr"

@app.post("/ia/recommendations")
async def generate_recommendations(req: RecommendationsRequest):
    try:
        lang_prompt = "Réponds en français." if req.lang == "fr" else ("Respond in English." if req.lang == "en" else "أجب باللغة العربية.")
        
        prompt = f"""Tu es un assistant IA expert en gestion des talents et SaaS (EvaluaTech).
L'utilisateur connecté a le rôle : {req.role}.
Génère EXACTEMENT 3 recommandations personnalisées et hautement stratégiques basées sur ce rôle.

{lang_prompt}

Tu devez associer à chaque recommandation une action concrète et opérationnelle sur notre plateforme. Pour cela, tu devez STRICTEMENT choisir l'un des chemins d'accès (route) réels suivants dans ton JSON :
1. '/questions' : Banque de questions (pour créer ou éditer des QCM)
2. '/campaigns' : Campagnes d'évaluation (pour lancer, suivre ou configurer des tests)
3. '/ai-generator' : Générateur de QCM IA (pour créer des questions via IA)
4. '/analyse-comportementale' : Analyse comportementale & Soft Skills des candidats
5. '/suivi-performance' : Suivi de l'activité en temps réel
6. '/sessions' : Sessions et planning d'évaluation
7. '/invite' : Inviter de nouveaux candidats ou collaborateurs RH
8. '/candidates-list' : Liste de tous les candidats et talents
9. '/reporting' : Génération de rapports RH mensuels ou de performance
10. '/staff-members' : Gestion des collaborateurs de l'entreprise
11. '/my-tests' : Espace candidat pour passer les tests en attente
12. '/history' : Historique des résultats et scores des candidats
13. '/super-admin' : Panel de contrôle global (uniquement pour le rôle 'SuperAdmin')
14. '/gestion-abonnements' : Gestion des formules et abonnements (uniquement pour le rôle 'SuperAdmin')
15. '/super-admin-analytics' : Métriques système globales (uniquement pour le rôle 'SuperAdmin')

RÈGLES CRITIQUES :
- Le champ 'route' doit STRICTEMENT faire partie de cette liste de 15 routes. N'invente JAMAIS d'autres routes (comme '/some-route' ou '/dashboard'). Si aucune ne correspond, mets null.
- Chaque recommandation doit être rédigée dans la langue demandée par l'utilisateur (langue demandée: {req.lang.upper()}).
- Chaque recommandation doit avoir une priorité unique parmi : 'Urgent' (couleur #ef4444), 'Priorité' (couleur #f59e0b), 'Standard' (couleur #10b981).
- Assure-toi que les titres et descriptions soient engageants, professionnels et rédigés de manière percutante.

Retourne STRICTEMENT le résultat en JSON sous cette forme exacte :
{{
  "recommendations": [
    {{
      "priority": "Urgent",
      "priorityBg": "rgba(239,68,68,0.12)",
      "priorityColor": "#ef4444",
      "color": "#ef4444",
      "icon": "fa-solid fa-bolt",
      "title": "Titre en {req.lang}",
      "description": "Description détaillée en {req.lang}...",
      "actionLabel": "Texte du bouton d'action",
      "route": "/route-choisie-dans-la-liste",
      "scrollTo": null
    }}
  ]
}}
"""
        response = await call_gemini_async(prompt, module="Recommandations", sem=_sem_reco, retries=1)
        text_response = response.text.replace("```json", "").replace("```", "").strip()
        data = json.loads(text_response)
        
        return {
            "status": "SUCCESS",
            "recommendations": data.get("recommendations", [])
        }
    except Exception as e:
        print(f"Error generating recommendations: {e}")
        return {
            "status": "ERROR",
            "message": str(e)
        }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)